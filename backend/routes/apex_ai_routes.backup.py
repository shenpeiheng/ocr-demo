"""
Apex AI 需求文档生成接口。

流程：
1. 用户在前端用自然语言描述业务需求（表单 / 报表）。
2. 后端调用 LLM 结构化输出 MD.050 所需字段。
3. 依据用户选择的模板类型（form / report），用 python-docx 生成 .docx。

参考：frontend/apex_ai 下的两份现有 MD.050 文档（表单/报表）为格式蓝本。
"""

from __future__ import annotations

import json
import re
from datetime import datetime
from json import JSONDecoder
from pathlib import Path
from typing import Any, Dict, Iterator, List
from uuid import uuid4

import requests
from flask import Blueprint, Response, jsonify, request, send_file, stream_with_context

from config import Config

apex_ai_bp = Blueprint("apex_ai", __name__)

ROOT_DIR = Path(__file__).resolve().parents[2]
APEX_AI_DIR = ROOT_DIR / "frontend" / "apex_ai"
EXPORTS_DIR = APEX_AI_DIR / "exports"

TEMPLATE_CATALOG: Dict[str, Dict[str, str]] = {
    "form": {
        "key": "form",
        "label": "表单类（MD.050）",
        "description": "适合业务表单类需求，如库存出入库、单据维护、审批流程等。",
        "reference": "MD.050-JUSENERP-ERP-INV-库存出入库平台V2.0(1).docx",
        "format": "docx",
    },
    "report": {
        "key": "report",
        "label": "报表类（MD.050）",
        "description": "适合报表类需求，如统计报表、明细报表、参数化查询报表等。",
        "reference": "YX&HB-ERP-MD050-HBCUX.离散任务缺料报表(1).docx",
        "format": "docx",
    },
}


def _resolve_model(model_alias: str = "") -> str:
    return Config.resolve_llm_model(model_alias)


def _normalize_modelscope_chat_url(model_key: str = "") -> str:
    base_url = Config.resolve_llm_url(model_key)
    if base_url.endswith("/chat/completions"):
        return base_url
    return f"{base_url.rstrip('/')}/chat/completions"


def _extract_json_object(raw_text: str) -> Dict[str, Any]:
    text = (raw_text or "").strip()
    if not text:
        return {}

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    fenced_match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.S)
    if fenced_match:
        try:
            return json.loads(fenced_match.group(1))
        except json.JSONDecodeError:
            pass

    brace_index = text.find("{")
    if brace_index >= 0:
        decoder = JSONDecoder()
        try:
            parsed, _ = decoder.raw_decode(text[brace_index:])
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass
    return {}


def _sse_event(event: str, data: Dict[str, Any]) -> str:
    return f"event: {event}\ndata: {json.dumps(data, ensure_ascii=False)}\n\n"


def _normalize_messages(messages: List[Dict[str, Any]]) -> List[Dict[str, str]]:
    normalized: List[Dict[str, str]] = []
    for item in (messages or [])[-12:]:
        role = item.get("role", "user")
        if role not in {"user", "assistant"}:
            continue
        content = str(item.get("content", "")).strip()
        if not content:
            continue
        normalized.append({"role": role, "content": content})
    return normalized


def _build_form_system_prompt() -> str:
    return """你是"MD.050 需求分析助手（表单类）"。

目标：把用户用自然语言描述的业务需求，结构化为 Oracle EBS 表单类 MD.050 完整文档所需的所有内容。

规则：
- 使用中文。
- 严格输出 JSON，不要输出 Markdown 代码块，不要输出额外说明。
- 界面数量根据用户描述判断，通常 2-6 个。
- 每个界面必须有完整的 fields（字段清单）、actions（按钮）、rules（业务规则）。
- 生成完整的 6 大块内容：文档控制、前言、概述、总体结构设计、详细功能设计、未决与已决问题。
- **重要**：即使用户描述简单，也要基于常见业务场景主动推理并补充完整的字段、按钮、规则，生成一个可用的完整初稿。
- **不要追问用户**：clarify_questions 留空数组，直接生成完整文档。用户如果不满意会主动反馈修改。
- phase 始终设为 "ready"，表示文档已就绪。

JSON schema:
{
  "phase": "clarify|ready",
  "document_title": "文档主标题，例如 采购订单管理系统",
  "project_name": "项目名，例如 XX集团 ERP 实施项目",
  "module_code": "模块代码，例如 ERP-PO",
  "version": "版本号，例如 V1.0",

  "section_3_overview": {
    "business_background": "3.1 业务背景说明，2-4 段完整描述当前业务痛点和系统目标",
    "business_requirements": ["3.2 相关业务需求要点1（完整句子）", "要点2", "要点3"],
    "problems_solved": ["3.3 开发后解决的问题1（完整句子）", "问题2", "问题3"],
    "usage_notes": {
      "target_users": "使用对象，例如 采购部、仓库、财务",
      "responsibility": "所属职责，例如 采购订单管理",
      "menu_path": "菜单位置，例如 采购管理 > 订单维护 > 订单录入",
      "prerequisite": "前提条件，例如 已完成供应商主数据维护"
    }
  },

  "section_4_structure": {
    "process_description": "4.1 相关业务流程说明，2-4 段完整描述业务流程的各个环节和数据流转",
    "screen_list_summary": "4.2 界面列表概述，1-2 段说明本功能包含哪些界面及其作用",
    "screen_overview": "4.3 界面概述，1-2 段说明各界面之间的关系和整体交互流程"
  },

  "section_5_screens": [
    {
      "name": "界面名称，例如 采购申请单维护",
      "type": "维护|查询|审批|执行|打印",
      "purpose": "界面用途，2-3 句完整说明",
      "fields": [
        {"name": "字段名", "type": "文本|日期|下拉|数字|复选|LOV", "required": "是|否", "description": "字段说明"}
      ],
      "actions": ["按钮1：功能说明", "按钮2：功能说明"],
      "rules": ["业务规则1（完整句子）", "业务规则2"],
      "screen_note": "界面示意说明，1-2 句描述界面布局或特殊说明"
    }
  ],

  "section_6_issues": {
    "open_issues": ["未决问题1（如果没有就留空数组）"],
    "closed_issues": ["已决问题1（如果没有就留空数组）"]
  },

  "reply": "给左侧聊天窗口的简短回复",
  "clarify_questions": ["还需要澄清的问题1", "问题2"]
}
"""


def _build_report_system_prompt() -> str:
    return """你是"MD.050 需求分析助手（报表类）"。

目标：把用户用自然语言描述的报表需求，结构化为 Oracle EBS 报表类 MD.050 所需字段。

规则：
- 使用中文。
- 严格输出 JSON，不要输出 Markdown 代码块，不要输出额外说明。
- 报表数量根据用户描述判断，通常 1-3 个。
- 每个报表必须有 parameters（查询参数）和 columns（输出列）。
- 若信息不足，仍然输出结构化 JSON，把不确定的字段留空字符串，并在 clarify_questions 里列出关键澄清问题。

JSON schema:
{
  "phase": "clarify|ready",
  "document_title": "报表文档主标题，例如 离散任务缺料报表",
  "project_name": "项目名",
  "module_code": "模块代码，例如 ERP-WIP",
  "version": "版本号，例如 V1.0",
  "business_background": "业务背景，2-4 句",
  "business_requirements": ["相关业务需求要点1"],
  "problems_solved": ["开发后解决的问题1"],
  "usage_notes": {
    "target_users": "使用对象",
    "responsibility": "所属职责",
    "menu_path": "菜单位置",
    "prerequisite": "前提条件"
  },
  "process_description": "取数逻辑与流程说明，1-3 段",
  "reports": [
    {
      "name": "报表名称",
      "purpose": "报表用途，1-2 句",
      "output_format": "Excel|PDF|页面展示",
      "trigger": "触发方式，例如 请求集手工提交/定时任务",
      "parameters": [
        {"name": "参数名", "type": "参数类型（日期/文本/LOV/数字）", "required": "是|否", "description": "参数说明", "default_value": "默认值"}
      ],
      "columns": [
        {"name": "列名", "type": "数据类型（文本/数字/日期）", "source": "取数来源，例如 库存事务表", "description": "列说明"}
      ],
      "sort_group": "排序与分组规则说明",
      "data_logic": "取数逻辑，2-4 句",
      "notes": ["特殊说明1", "说明2"]
    }
  ],
  "reply": "给左侧聊天窗口的简短回复",
  "clarify_questions": ["还需要澄清的问题1"]
}
"""


def _system_prompt_for(template_key: str) -> str:
    if template_key == "report":
        return _build_report_system_prompt()
    return _build_form_system_prompt()


def _request_completion(
    messages: List[Dict[str, str]],
    max_tokens: int,
    timeout_seconds: int,
    model: str = "",
    stream: bool = False,
) -> Any:
    api_key = Config.resolve_llm_key(model)
    payload = {
        "model": _resolve_model(model),
        "messages": messages,
        "temperature": 0.2,
        "max_tokens": max_tokens,
    }
    if stream:
        payload["stream"] = True

    try:
        response = requests.post(
            _normalize_modelscope_chat_url(model),
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json=payload,
            timeout=timeout_seconds,
            stream=stream,
        )
    except requests.exceptions.ReadTimeout as exc:
        raise RuntimeError(f"模型服务响应超时（{timeout_seconds}秒），请缩小本次生成范围或重试。") from exc
    except requests.exceptions.RequestException as exc:
        raise RuntimeError(f"模型服务网络请求失败: {exc}") from exc

    if response.status_code != 200:
        text = response.text[:1000] if not stream else ""
        try:
            response.close()
        except Exception:
            pass
        raise RuntimeError(f"模型服务调用失败: {response.status_code} {text}")

    if stream:
        return response

    data = response.json()
    choice = (data.get("choices") or [{}])[0]
    message = choice.get("message") or {}
    return {
        "content": message.get("content", "") or "",
        "finish_reason": choice.get("finish_reason", "") or "",
    }


def _iter_stream(response: requests.Response) -> Iterator[Dict[str, Any]]:
    try:
        response.encoding = "utf-8"
        for raw_line in response.iter_lines(decode_unicode=False):
            if not raw_line:
                continue
            line = raw_line.decode("utf-8", errors="ignore") if isinstance(raw_line, bytes) else str(raw_line)
            if not line.startswith("data:"):
                continue
            payload_text = line[5:].strip()
            if not payload_text or payload_text == "[DONE]":
                continue
            try:
                payload = json.loads(payload_text)
            except json.JSONDecodeError:
                continue
            choice = (payload.get("choices") or [{}])[0]
            delta = choice.get("delta") or {}
            message = choice.get("message") or {}
            finish_reason = choice.get("finish_reason", "") or ""
            content = delta.get("content")
            if content is None:
                content = message.get("content", "") or ""
            yield {
                "content": str(content or ""),
                "finish_reason": finish_reason,
            }
    finally:
        try:
            response.close()
        except Exception:
            pass


# ==================== 分析结果规范化 ====================


def _as_list_of_str(value: Any) -> List[str]:
    if isinstance(value, list):
        return [str(item).strip() for item in value if str(item).strip()]
    text = str(value or "").strip()
    if not text:
        return []
    lines = [line.strip("-•· \t") for line in text.splitlines() if line.strip()]
    return lines or [text]


def _as_dict(value: Any) -> Dict[str, str]:
    if isinstance(value, dict):
        return {str(k): str(v).strip() for k, v in value.items() if str(v).strip()}
    return {}


def _normalize_screen(item: Any) -> Dict[str, Any]:
    if not isinstance(item, dict):
        return {}
    fields = []
    for field in item.get("fields") or []:
        if not isinstance(field, dict):
            continue
        fields.append({
            "name": str(field.get("name", "")).strip(),
            "type": str(field.get("type", "")).strip(),
            "required": str(field.get("required", "")).strip() or "否",
            "description": str(field.get("description", "")).strip(),
        })
    return {
        "name": str(item.get("name", "")).strip() or "未命名界面",
        "type": str(item.get("type", "")).strip() or "维护",
        "purpose": str(item.get("purpose", "")).strip(),
        "fields": fields,
        "actions": _as_list_of_str(item.get("actions")),
        "rules": _as_list_of_str(item.get("rules")),
    }


def _normalize_report(item: Any) -> Dict[str, Any]:
    if not isinstance(item, dict):
        return {}
    parameters = []
    for param in item.get("parameters") or []:
        if not isinstance(param, dict):
            continue
        parameters.append({
            "name": str(param.get("name", "")).strip(),
            "type": str(param.get("type", "")).strip(),
            "required": str(param.get("required", "")).strip() or "否",
            "description": str(param.get("description", "")).strip(),
            "default_value": str(param.get("default_value", "")).strip(),
        })
    columns = []
    for column in item.get("columns") or []:
        if not isinstance(column, dict):
            continue
        columns.append({
            "name": str(column.get("name", "")).strip(),
            "type": str(column.get("type", "")).strip(),
            "source": str(column.get("source", "")).strip(),
            "description": str(column.get("description", "")).strip(),
        })
    return {
        "name": str(item.get("name", "")).strip() or "未命名报表",
        "purpose": str(item.get("purpose", "")).strip(),
        "output_format": str(item.get("output_format", "")).strip() or "Excel",
        "trigger": str(item.get("trigger", "")).strip(),
        "parameters": parameters,
        "columns": columns,
        "sort_group": str(item.get("sort_group", "")).strip(),
        "data_logic": str(item.get("data_logic", "")).strip(),
        "notes": _as_list_of_str(item.get("notes")),
    }


def _normalize_analysis(parsed: Dict[str, Any], template_key: str) -> Dict[str, Any]:
    if not isinstance(parsed, dict):
        parsed = {}

    # 支持新旧两种 JSON 结构
    section_3 = parsed.get("section_3_overview", {})
    section_4 = parsed.get("section_4_structure", {})
    section_5 = parsed.get("section_5_screens", [])
    section_6 = parsed.get("section_6_issues", {})

    # 兼容旧结构：如果新字段为空，从旧字段读取
    if not section_3:
        section_3 = {
            "business_background": parsed.get("business_background", ""),
            "business_requirements": parsed.get("business_requirements", []),
            "problems_solved": parsed.get("problems_solved", []),
            "usage_notes": parsed.get("usage_notes", {}),
        }
    if not section_4:
        section_4 = {
            "process_description": parsed.get("process_description", ""),
            "screen_list_summary": "",
            "screen_overview": "",
        }
    if not section_5:
        section_5 = parsed.get("screens", [])

    normalized = {
        "template_key": template_key,
        "phase": str(parsed.get("phase", "") or "ready").strip(),
        "document_title": str(parsed.get("document_title", "") or "").strip() or "需求说明文档",
        "project_name": str(parsed.get("project_name", "") or "").strip(),
        "module_code": str(parsed.get("module_code", "") or "").strip(),
        "version": str(parsed.get("version", "") or "V1.0").strip() or "V1.0",
        "reply": str(parsed.get("reply", "") or "").strip(),
        "clarify_questions": _as_list_of_str(parsed.get("clarify_questions")),
        # 新结构
        "section_3_overview": section_3,
        "section_4_structure": section_4,
        "section_6_issues": section_6,
        # 兼容旧字段（用于前端展示）
        "business_background": str(section_3.get("business_background", "") or "").strip(),
        "business_requirements": _as_list_of_str(section_3.get("business_requirements")),
        "problems_solved": _as_list_of_str(section_3.get("problems_solved")),
        "usage_notes": _as_dict(section_3.get("usage_notes")),
        "process_description": str(section_4.get("process_description", "") or "").strip(),
    }

    if template_key == "report":
        normalized["reports"] = [
            item for item in (_normalize_report(r) for r in parsed.get("reports") or []) if item
        ]
    else:
        # 表单类：从新结构读取 screens
        normalized["section_5_screens"] = [
            item for item in (_normalize_screen(s) for s in section_5) if item
        ]
        # 兼容旧字段（用于前端展示）
        normalized["screens"] = normalized["section_5_screens"]

    return normalized


# ==================== docx 文档生成 ====================


def _configure_default_style(doc) -> None:
    from docx.shared import Pt

    style = doc.styles["Normal"]
    style.font.name = "微软雅黑"
    style.font.size = Pt(10.5)
    from docx.oxml.ns import qn
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _add_heading(doc, text: str, level: int) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    heading = doc.add_paragraph()
    heading.paragraph_format.space_before = Pt(10)
    heading.paragraph_format.space_after = Pt(6)
    run = heading.add_run(text)
    run.font.name = "微软雅黑"
    run.font.bold = True
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    else:
        run.font.size = Pt(11)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def _add_paragraph(doc, text: str, bold: bool = False, italic: bool = False) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = "微软雅黑"
    run.font.size = Pt(10.5)
    run.font.bold = bold
    run.font.italic = italic
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _add_bullet(doc, text: str) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    para = doc.add_paragraph(style="List Bullet")
    para.paragraph_format.space_after = Pt(2)
    run = para.runs[0] if para.runs else para.add_run("")
    if not para.runs:
        run = para.add_run(text)
    else:
        run.text = text
    run.font.name = "微软雅黑"
    run.font.size = Pt(10.5)
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _set_cell_text(cell, text: str, bold: bool = False) -> None:
    from docx.shared import Pt
    from docx.oxml.ns import qn

    cell.text = ""
    para = cell.paragraphs[0]
    run = para.add_run(text if text else "")
    run.font.name = "微软雅黑"
    run.font.size = Pt(10)
    run.font.bold = bold
    run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")


def _shade_cell(cell, hex_color: str) -> None:
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    cell._tc.get_or_add_tcPr().append(shd)


def _add_table(doc, headers: List[str], rows: List[List[str]], first_col_bold: bool = False):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        _set_cell_text(cell, header, bold=True)
        _shade_cell(cell, "D9E2F3")
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        for c_idx, value in enumerate(row_data):
            if c_idx >= len(row.cells):
                break
            _set_cell_text(row.cells[c_idx], value, bold=(first_col_bold and c_idx == 0))
    return table


def _add_cover_page(doc, analysis: Dict[str, Any], template_key: str) -> None:
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    project_name = analysis.get("project_name") or "ERP 实施项目"
    document_title = analysis.get("document_title") or "需求说明文档"
    version = analysis.get("version") or "V1.0"
    doc_type_label = "报表类需求说明文档（MD.050）" if template_key == "report" else "表单类需求说明文档（MD.050）"

    def _centered(text: str, size: int, bold: bool = False, color=None) -> None:
        para = doc.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        para.paragraph_format.space_after = Pt(6)
        run = para.add_run(text)
        run.font.name = "微软雅黑"
        run.font.size = Pt(size)
        run.font.bold = bold
        if color is not None:
            run.font.color.rgb = color
        run.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")

    for _ in range(3):
        doc.add_paragraph()

    _centered(project_name, 20, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    doc.add_paragraph()
    _centered(document_title, 26, bold=True, color=RGBColor(0xAD, 0x06, 0x3E))
    doc.add_paragraph()
    _centered(doc_type_label, 14, bold=False)

    for _ in range(4):
        doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = "Table Grid"
    info_rows = [
        ("作者", "AI 自动生成"),
        ("创建日期", datetime.now().strftime("%Y-%m-%d")),
        ("版本", version),
        ("模块代码", analysis.get("module_code") or "-"),
    ]
    for i, (label, value) in enumerate(info_rows):
        _set_cell_text(info_table.rows[i].cells[0], label, bold=True)
        _shade_cell(info_table.rows[i].cells[0], "F2F2F2")
        _set_cell_text(info_table.rows[i].cells[1], value)

    doc.add_page_break()


def _add_document_control(doc, analysis: Dict[str, Any]) -> None:
    _add_heading(doc, "1. 文档控制", level=1)

    _add_heading(doc, "1.1 更改记录", level=2)
    _add_table(
        doc,
        headers=["日期", "姓名", "版本", "变更参考"],
        rows=[
            [datetime.now().strftime("%Y-%m-%d"), "AI 自动生成", analysis.get("version") or "V1.0", "基于自然语言需求初次生成"],
        ],
    )

    _add_heading(doc, "1.2 审核", level=2)
    _add_table(
        doc,
        headers=["姓名", "职位"],
        rows=[["", ""], ["", ""], ["", ""]],
    )

    _add_heading(doc, "1.3 分发", level=2)
    _add_table(
        doc,
        headers=["编号", "名称", "地点"],
        rows=[["1", "", ""], ["2", "", ""]],
    )


def _add_preface(doc, analysis: Dict[str, Any]) -> None:
    _add_heading(doc, "2. 前言", level=1)
    _add_heading(doc, "2.1 说明", level=2)
    _add_paragraph(doc, "本文档描述了基于自然语言需求梳理出的功能设计说明，供项目开发、测试和用户参考。")
    _add_paragraph(doc, "如客户对本文档有意见或反馈，请及时提出，以便进一步修订与完善。")

    _add_heading(doc, "2.2 参与人员", level=2)
    _add_paragraph(doc, "关键用户、业务顾问、开发顾问（按项目实际情况填写）。")


def _add_overview(doc, analysis: Dict[str, Any]) -> None:
    _add_heading(doc, "3. 概述", level=1)

    _add_heading(doc, "3.1 业务背景", level=2)
    background = analysis.get("business_background") or "（暂无描述，请后续补充）"
    for paragraph in background.split("\n"):
        text = paragraph.strip()
        if text:
            _add_paragraph(doc, text)

    _add_heading(doc, "3.2 相关业务需求", level=2)
    requirements = analysis.get("business_requirements") or []
    if requirements:
        for item in requirements:
            _add_bullet(doc, item)
    else:
        _add_paragraph(doc, "（暂无描述，请后续补充）")

    _add_heading(doc, "3.3 开发后解决的问题", level=2)
    problems = analysis.get("problems_solved") or []
    if problems:
        for item in problems:
            _add_bullet(doc, item)
    else:
        _add_paragraph(doc, "（暂无描述，请后续补充）")

    _add_heading(doc, "3.4 使用说明", level=2)
    usage = analysis.get("usage_notes") or {}
    _add_table(
        doc,
        headers=["项", "说明"],
        rows=[
            ["使用对象", usage.get("target_users", "") or ""],
            ["所属职责", usage.get("responsibility", "") or ""],
            ["菜单位置", usage.get("menu_path", "") or ""],
            ["前提条件", usage.get("prerequisite", "") or ""],
        ],
        first_col_bold=True,
    )


def _add_structure_design_form(doc, analysis: Dict[str, Any]) -> None:
    _add_heading(doc, "4. 总体结构设计", level=1)
    _add_heading(doc, "4.1 相关业务流程说明", level=2)
    process = analysis.get("process_description") or "（暂无描述，请后续补充）"
    for paragraph in process.split("\n"):
        text = paragraph.strip()
        if text:
            _add_paragraph(doc, text)

    _add_heading(doc, "4.2 界面列表", level=2)
    screens = analysis.get("screens") or []
    if screens:
        _add_table(
            doc,
            headers=["界面序号", "初始界面", "界面名称", "界面类型", "主从关系", "主界面号", "说明"],
            rows=[
                [
                    str(i + 1),
                    "",  # 初始界面
                    screen["name"],
                    screen.get("type", ""),
                    "",  # 主从关系
                    "",  # 主界面号
                    screen.get("purpose", "")
                ]
                for i, screen in enumerate(screens)
            ],
        )
    else:
        _add_paragraph(doc, "（暂无界面，请补充需求）")

    _add_heading(doc, "4.3 流程设计", level=2)
    _add_paragraph(doc, "（请根据实际业务流程补充流程图或流程说明）")


def _add_detail_form(doc, analysis: Dict[str, Any]) -> None:
    _add_heading(doc, "5. 详细功能设计", level=1)
    screens = analysis.get("screens") or []
    if not screens:
        _add_paragraph(doc, "（暂无界面明细，请补充需求）")
        return

    for i, screen in enumerate(screens):
        _add_heading(doc, f"5.{i + 1} 界面{i + 1}-{screen['name']}", level=2)

        _add_heading(doc, "界面设计", level=3)

        _add_heading(doc, "界面用途", level=3)
        _add_paragraph(doc, screen.get("purpose") or "（未描述）")

        fields = screen.get("fields") or []
        _add_heading(doc, "字段清单", level=3)
        if fields:
            _add_table(
                doc,
                headers=["字段名", "类型", "是否必填", "说明"],
                rows=[[f.get("name", ""), f.get("type", ""), f.get("required", ""), f.get("description", "")] for f in fields],
            )
        else:
            _add_paragraph(doc, "（无字段定义）")

        actions = screen.get("actions") or []
        _add_heading(doc, "按钮与操作", level=3)
        if actions:
            for action in actions:
                _add_bullet(doc, action)
        else:
            _add_paragraph(doc, "（未定义按钮）")

        rules = screen.get("rules") or []
        _add_heading(doc, "业务规则", level=3)
        if rules:
            for rule in rules:
                _add_bullet(doc, rule)
        else:
            _add_paragraph(doc, "（未定义业务规则）")


def _add_structure_design_report(doc, analysis: Dict[str, Any]) -> None:
    _add_heading(doc, "4. 总体结构设计", level=1)
    _add_heading(doc, "4.1 取数逻辑与流程说明", level=2)
    process = analysis.get("process_description") or "（暂无描述，请后续补充）"
    for paragraph in process.split("\n"):
        text = paragraph.strip()
        if text:
            _add_paragraph(doc, text)

    _add_heading(doc, "4.2 报表列表", level=2)
    reports = analysis.get("reports") or []
    if reports:
        _add_table(
            doc,
            headers=["序号", "报表名称", "输出格式", "触发方式", "说明"],
            rows=[
                [str(i + 1), r["name"], r.get("output_format", ""), r.get("trigger", ""), r.get("purpose", "")]
                for i, r in enumerate(reports)
            ],
        )
    else:
        _add_paragraph(doc, "（暂无报表，请补充需求）")


def _add_detail_report(doc, analysis: Dict[str, Any]) -> None:
    _add_heading(doc, "5. 详细报表设计", level=1)
    reports = analysis.get("reports") or []
    if not reports:
        _add_paragraph(doc, "（暂无报表明细，请补充需求）")
        return

    for i, report in enumerate(reports):
        _add_heading(doc, f"5.{i + 1} 报表{i + 1}-{report['name']}", level=2)

        _add_heading(doc, "报表用途", level=3)
        _add_paragraph(doc, report.get("purpose") or "（未描述）")

        params = report.get("parameters") or []
        _add_heading(doc, "报表参数", level=3)
        if params:
            _add_table(
                doc,
                headers=["参数名", "类型", "是否必填", "默认值", "说明"],
                rows=[[p.get("name", ""), p.get("type", ""), p.get("required", ""), p.get("default_value", ""), p.get("description", "")] for p in params],
            )
        else:
            _add_paragraph(doc, "（无参数定义）")

        columns = report.get("columns") or []
        _add_heading(doc, "报表输出列", level=3)
        if columns:
            _add_table(
                doc,
                headers=["列名", "数据类型", "取数来源", "说明"],
                rows=[[c.get("name", ""), c.get("type", ""), c.get("source", ""), c.get("description", "")] for c in columns],
            )
        else:
            _add_paragraph(doc, "（无输出列定义）")

        _add_heading(doc, "取数逻辑", level=3)
        _add_paragraph(doc, report.get("data_logic") or "（未描述）")

        sort_group = report.get("sort_group") or ""
        if sort_group:
            _add_heading(doc, "排序与分组", level=3)
            _add_paragraph(doc, sort_group)

        notes = report.get("notes") or []
        if notes:
            _add_heading(doc, "特殊说明", level=3)
            for note in notes:
                _add_bullet(doc, note)


def _replace_text_in_para(paragraph, old_text: str, new_text: str) -> bool:
    """替换段落中的文本，保留格式"""
    if old_text in paragraph.text:
        for run in paragraph.runs:
            if old_text in run.text:
                run.text = run.text.replace(old_text, new_text)
        return True
    return False


def _update_table_row(table, row_index: int, values: list):
    """更新表格指定行的内容"""
    if row_index < len(table.rows):
        row = table.rows[row_index]
        for col_index, value in enumerate(values):
            if col_index < len(row.cells):
                cell = row.cells[col_index]
                cell.text = str(value) if value else ""


def _clear_table_data_rows(table, start_row: int = 1):
    """清空表格数据行（保留表头）"""
    for i in range(start_row, len(table.rows)):
        for cell in table.rows[i].cells:
            cell.text = ""


def _add_open_closed_issues(doc) -> None:
    _add_heading(doc, "6. 未结与已结问题", level=1)
    _add_heading(doc, "6.1 未结问题", level=2)
    _add_table(
        doc,
        headers=["序号", "问题", "解决方案", "负责人", "目标日期", "实际日期"],
        rows=[["", "", "", "", "", ""] for _ in range(5)],
    )
    _add_heading(doc, "6.2 已结问题", level=2)
    _add_table(
        doc,
        headers=["序号", "问题", "解决方案", "负责人", "目标日期", "实际日期"],
        rows=[["", "", "", "", "", ""] for _ in range(3)],
    )


def _build_docx(analysis: Dict[str, Any], template_key: str) -> "Path":
    from docx import Document
    import shutil
    import tempfile

    # 加载对应的模板文件
    template_filename = TEMPLATE_CATALOG.get(template_key, {}).get("reference", "")
    template_format = TEMPLATE_CATALOG.get(template_key, {}).get("format", "docx")

    if not template_filename:
        raise RuntimeError(f"未找到模板类型 {template_key} 的参考文档")

    template_path = APEX_AI_DIR / template_filename
    if not template_path.exists():
        raise RuntimeError(f"模板文件不存在: {template_path}")

    # 如果是 .doc 格式，提示用户需要转换
    if template_format == "doc":
        raise RuntimeError(
            f"报表类模板是旧版 .doc 格式，暂不支持。\n"
            f"请手动将模板转换为 .docx 格式：\n"
            f"1. 用 Word 打开 {template_filename}\n"
            f"2. 另存为 .docx 格式\n"
            f"3. 将新文件放在 frontend/apex_ai/ 目录下"
        )

    # 准备输出路径
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    slug = _slugify(analysis.get("document_title") or "requirement")
    export_id = datetime.now().strftime("%Y%m%d_%H%M%S") + "_" + uuid4().hex[:6]
    docx_name = f"{export_id}_{slug}.docx"
    docx_path = EXPORTS_DIR / docx_name

    # 使用临时文件避免权限问题
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp:
        tmp_path = Path(tmp.name)
        # 复制模板到临时文件
        shutil.copy2(str(template_path), str(tmp_path))

        # 打开临时文件进行内容替换
        doc = Document(str(tmp_path))

        # 提取数据
        document_title = analysis.get("document_title", "需求说明文档")
        project_name = analysis.get("project_name", "ERP实施项目")
        version = analysis.get("version", "V1.0")

        if template_key == "form":
            # 表单类模板替换逻辑
            # 1. 替换标题
            for para in doc.paragraphs:
                if "库存出入库平台" in para.text:
                    _replace_text_in_para(para, "库存出入库平台", document_title)
                if "ERP实施项目" in para.text:
                    _replace_text_in_para(para, "ERP实施项目", project_name)

            # 2. 更新表格
            if len(doc.tables) >= 8:
                # 表格1: 更改记录表
                _update_table_row(doc.tables[1], 1, [
                    datetime.now().strftime("%Y-%m-%d"),
                    "AI 自动生成",
                    version,
                    "基于自然语言需求初次生成"
                ])

                # 表格4: 使用说明表
                usage_notes = analysis.get("usage_notes", {})
                _update_table_row(doc.tables[4], 1, ["使用对象", usage_notes.get("target_users", "")])
                _update_table_row(doc.tables[4], 2, ["所属职责", usage_notes.get("responsibility", "")])
                _update_table_row(doc.tables[4], 3, ["菜单位置", usage_notes.get("menu_path", "")])
                _update_table_row(doc.tables[4], 4, ["前提条件", usage_notes.get("prerequisite", "")])

                # 表格5: 界面列表表
                screens = analysis.get("screens", [])
                if screens:
                    _clear_table_data_rows(doc.tables[5], start_row=1)
                    for i, screen in enumerate(screens[:6]):
                        _update_table_row(doc.tables[5], i + 1, [
                            str(i + 1),
                            "",  # 初始界面
                            screen.get("name", ""),
                            screen.get("type", ""),
                            "",  # 主从关系
                            "",  # 主界面号
                            screen.get("purpose", "")
                        ])

            # 3. 替换内容章节
            _replace_form_content_sections(doc, analysis)

        elif template_key == "report":
            # 报表类模板替换逻辑
            # 替换段落中的关键文本
            for para in doc.paragraphs:
                if "离散任务缺料报表" in para.text:
                    _replace_text_in_para(para, "离散任务缺料报表", document_title)
                if "YX&HB-ERP-MD050-WIP" in para.text:
                    _replace_text_in_para(para, "YX&HB-ERP-MD050-WIP", analysis.get("module_code", ""))
                if "阳信&汉邦ERP项目" in para.text:
                    _replace_text_in_para(para, "阳信&汉邦ERP项目", project_name)

            # 更新表格内容
            # 表格1: 更改记录表 (8行 x 4列)
            if len(doc.tables) > 1:
                _update_table_row(doc.tables[1], 1, [
                    datetime.now().strftime("%Y-%m-%d"),
                    "AI 自动生成",
                    version,
                    "基于自然语言需求初次生成"
                ])

            # 表格4: 使用说明表 (25行 x 4列，实际是键值对形式)
            usage_notes = analysis.get("usage_notes", {})
            if len(doc.tables) > 4:
                # 找到"使用对象"、"所属职责"等行并更新
                for row in doc.tables[4].rows[1:]:
                    key = row.cells[1].text.strip()
                    if "使用对象" in key:
                        row.cells[2].text = usage_notes.get("target_users", "")
                    elif "所属职责" in key:
                        row.cells[2].text = usage_notes.get("responsibility", "")
                    elif "菜单位置" in key:
                        row.cells[2].text = usage_notes.get("menu_path", "")
                    elif "前提条件" in key:
                        row.cells[2].text = usage_notes.get("prerequisite", "")

        # 保存到最终位置
        doc.save(str(docx_path))

        # 清理临时文件
        try:
            tmp_path.unlink()
        except Exception:
            pass

    meta = {
        "docx_name": docx_name,
        "export_id": export_id,
        "title": analysis.get("document_title") or "需求说明文档",
        "template_key": template_key,
        "template_label": TEMPLATE_CATALOG.get(template_key, {}).get("label", ""),
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "version": analysis.get("version") or "V1.0",
    }
    meta_path = docx_path.with_suffix(".json")
    meta_path.write_text(json.dumps(meta, ensure_ascii=False, indent=2), encoding="utf-8")

    return docx_path


def _replace_form_content_sections(doc: "Document", analysis: Dict[str, Any]) -> None:
    """替换表单类模板的内容章节 - 完整的6大块"""
    from docx.shared import Pt

    # ===== 3. 概述 =====
    section_3 = analysis.get("section_3_overview", {})

    # 3.1 业务背景
    business_bg = section_3.get("business_background", "")
    if business_bg:
        _replace_section_content(doc, "业务背景", business_bg)

    # 3.2 相关业务需求
    business_requirements = section_3.get("business_requirements", [])
    if business_requirements:
        req_text = "\n".join([f"{i+1}. {req}" for i, req in enumerate(business_requirements)])
        _replace_section_content(doc, "相关业务需求", req_text)

    # 3.3 开发后解决的问题
    problems_solved = section_3.get("problems_solved", [])
    if problems_solved:
        problems_text = "\n".join([f"{i+1}. {p}" for i, p in enumerate(problems_solved)])
        _replace_section_content(doc, "开发后解决的问题", problems_text)

    # 3.4 使用说明 - 段落内容（表格由 _update_table_row 处理）
    usage_notes = section_3.get("usage_notes", {})
    usage_text = f"""本功能的使用说明如下：
- 使用对象：{usage_notes.get('target_users', '')}
- 所属职责：{usage_notes.get('responsibility', '')}
- 菜单路径：{usage_notes.get('menu_path', '')}
- 前提条件：{usage_notes.get('prerequisite', '')}"""
    _replace_section_content(doc, "使用说明", usage_text)

    # ===== 4. 总体结构设计 =====
    section_4 = analysis.get("section_4_structure", {})

    # 4.1 相关业务流程说明
    process_desc = section_4.get("process_description", "")
    if process_desc:
        _replace_section_content(doc, "相关业务流程说明", process_desc)

    # 4.2 界面列表 - 段落内容（表格由前面的代码处理）
    screen_list_summary = section_4.get("screen_list_summary", "")
    if not screen_list_summary:
        screens = analysis.get("section_5_screens", [])
        if screens:
            screen_list_summary = f"本功能模块包含 {len(screens)} 个界面，分别为：\n"
            screen_list_summary += "\n".join([f"{i+1}. {s.get('name', '')} - {s.get('purpose', '')}"
                                             for i, s in enumerate(screens)])
    if screen_list_summary:
        _replace_section_content(doc, "界面列表", screen_list_summary)

    # 4.3 界面概述
    screen_overview = section_4.get("screen_overview", "")
    if not screen_overview:
        screens = analysis.get("section_5_screens", [])
        if screens:
            screen_overview = "各界面功能关系如下：\n"
            screen_overview += "\n".join([f"• {s.get('name', '')}（{s.get('type', '')}）：{s.get('purpose', '')}"
                                         for s in screens])
    if screen_overview:
        _replace_section_content(doc, "界面概述", screen_overview)

    # ===== 5. 详细功能设计 - 每个界面 =====
    screens = analysis.get("section_5_screens", [])
    for i, screen in enumerate(screens):
        screen_name = screen.get("name", f"界面{i+1}")
        screen_content = _build_screen_content_full(screen)
        _replace_screen_content(doc, screen_name, screen_content, i)

    # ===== 6. 未决与已决问题 =====
    section_6 = analysis.get("section_6_issues", {})

    # 6.1 未决问题
    open_issues = section_6.get("open_issues", [])
    if open_issues:
        open_text = "\n".join([f"{i+1}. {issue}" for i, issue in enumerate(open_issues)])
    else:
        open_text = "暂无未决问题。"
    _replace_section_content(doc, "未决问题", open_text)

    # 6.2 已决问题
    closed_issues = section_6.get("closed_issues", [])
    if closed_issues:
        closed_text = "\n".join([f"{i+1}. {issue}" for i, issue in enumerate(closed_issues)])
    else:
        closed_text = "暂无已决问题。"
    _replace_section_content(doc, "已决问题", closed_text)


def _replace_section_content(doc: "Document", title_keyword: str, new_content: str) -> None:
    """通过章节标题关键词查找并替换内容（只替换内容段落，不动标题）"""
    from docx.shared import Pt

    if not new_content:
        return

    paragraphs = doc.paragraphs
    title_idx = -1

    # 查找包含关键词的 SIM标题2 段落
    for i, para in enumerate(paragraphs):
        if (title_keyword in para.text and
            para.style.name in ['SIM标题2', 'SIM标题1'] and
            len(para.text.strip()) < 40):
            title_idx = i
            break

    if title_idx == -1:
        return

    # 找到下一个 SIM标题1 或 SIM标题2 的位置
    next_title_idx = len(paragraphs)
    for i in range(title_idx + 1, len(paragraphs)):
        if paragraphs[i].style.name in ['SIM标题1', 'SIM标题2']:
            next_title_idx = i
            break

    # 清空标题和下一个标题之间的所有 Body Text 段落
    for i in range(next_title_idx - 1, title_idx, -1):
        if i < len(paragraphs) and i > title_idx:
            para = paragraphs[i]
            # 只清空 Body Text，不清空标题
            if para.style.name not in ['SIM标题1', 'SIM标题2', 'SIM标题3']:
                para.text = ""

    # 在标题后的第一个段落插入新内容
    if title_idx + 1 < len(paragraphs):
        target_para = paragraphs[title_idx + 1]
        target_para.text = new_content
        # 设置字体
        for run in target_para.runs:
            run.font.name = "微软雅黑"
            run.font.size = Pt(10.5)


def _replace_screen_content(doc: "Document", screen_name: str, new_content: str, screen_index: int) -> None:
    """替换界面详细说明章节（查找 "界面X-名称" 格式的标题）"""
    from docx.shared import Pt

    if not new_content:
        return

    paragraphs = doc.paragraphs
    title_idx = -1

    # 查找界面标题（例如："界面1-采购申请单维护"）
    for i, para in enumerate(paragraphs):
        if (f"界面{screen_index+1}" in para.text and
            screen_name in para.text and
            para.style.name == 'SIM标题2'):
            title_idx = i
            break

    # 如果找不到精确匹配，尝试只匹配界面名称
    if title_idx == -1:
        for i, para in enumerate(paragraphs):
            if (screen_name in para.text and
                '界面' in para.text and
                para.style.name == 'SIM标题2'):
                title_idx = i
                break

    if title_idx == -1:
        return

    # 找到下一个界面标题的位置
    next_title_idx = len(paragraphs)
    for i in range(title_idx + 1, len(paragraphs)):
        para = paragraphs[i]
        # 下一个界面标题或大章节标题
        if (para.style.name == 'SIM标题2' and '界面' in para.text) or para.style.name == 'SIM标题1':
            next_title_idx = i
            break

    # 清空该界面的所有内容段落（保留标题和子标题）
    for i in range(next_title_idx - 1, title_idx, -1):
        if i < len(paragraphs) and i > title_idx:
            para = paragraphs[i]
            # 只清空 Body Text，保留所有标题
            if para.style.name not in ['SIM标题1', 'SIM标题2', 'SIM标题3']:
                para.text = ""

    # 在标题后的第一个段落插入新内容
    if title_idx + 1 < len(paragraphs):
        target_para = paragraphs[title_idx + 1]
        target_para.text = new_content
        # 设置字体
        for run in target_para.runs:
            run.font.name = "微软雅黑"
            run.font.size = Pt(10.5)


def _build_screen_content_full(screen: Dict[str, Any]) -> str:
    """构建完整的界面详细说明内容"""
    content_parts = []

    # 界面设计说明
    if screen.get("purpose"):
        content_parts.append(f"【界面用途】\n{screen['purpose']}\n")

    # 字段清单
    fields = screen.get("fields", [])
    if fields:
        content_parts.append("【字段清单】")
        for j, field in enumerate(fields, 1):
            required_mark = " [必填]" if field.get('required') == '是' else ""
            field_line = f"{j}. {field.get('name', '')}（{field.get('type', '')}）{required_mark}"
            if field.get('description'):
                field_line += f" - {field['description']}"
            content_parts.append(field_line)
        content_parts.append("")

    # 按钮功能
    actions = screen.get("actions", [])
    if actions:
        content_parts.append("【按钮功能】")
        for j, action in enumerate(actions, 1):
            content_parts.append(f"{j}. {action}")
        content_parts.append("")

    # 业务规则
    rules = screen.get("rules", [])
    if rules:
        content_parts.append("【业务规则】")
        for j, rule in enumerate(rules, 1):
            content_parts.append(f"{j}. {rule}")
        content_parts.append("")

    # 界面示意
    if screen.get("screen_note"):
        content_parts.append(f"【界面示意】\n{screen['screen_note']}")
    else:
        field_count = len(fields)
        action_count = len(actions)
        content_parts.append(f"【界面示意】\n本界面包含 {field_count} 个字段和 {action_count} 个操作按钮，采用标准表单布局。")

    return "\n".join(content_parts)


def _delete_paragraph(paragraph) -> None:
    """删除段落"""
    p = paragraph._element
    parent = p.getparent()
    if parent is not None:
        parent.remove(p)


def _insert_paragraph_after(paragraph, text: str) -> None:
    """在指定段落后插入新段落"""
    from docx.oxml import OxmlElement
    from docx.shared import Pt

    new_p = paragraph._element.getparent().insert(
        paragraph._element.getparent().index(paragraph._element) + 1,
        paragraph._element.__class__()
    )
    # 复制原段落的样式
    new_para = paragraph._p.__class__(new_p, paragraph._parent)
    new_para.text = text
    new_para.style = paragraph.style

    # 设置字体
    for run in new_para.runs:
        run.font.name = "微软雅黑"
        run.font.size = Pt(10.5)


def _slugify(text: str) -> str:
    value = re.sub(r"[^\w一-鿿-]+", "-", (text or "").strip(), flags=re.U)
    value = re.sub(r"-{2,}", "-", value).strip("-")
    return value[:48] or "requirement"


# ==================== 路由 ====================


@apex_ai_bp.route("/api/apex-ai/templates", methods=["GET"])
def get_templates():
    return jsonify({
        "success": True,
        "templates": list(TEMPLATE_CATALOG.values()),
        "default_template": "form",
    })


def _finalize_analysis(content: str, template_key: str) -> Dict[str, Any]:
    parsed = _extract_json_object(content)
    analysis = _normalize_analysis(parsed, template_key)
    return analysis


@apex_ai_bp.route("/api/apex-ai/analyze", methods=["POST"])
def analyze():
    data = request.get_json(silent=True) or {}
    messages = data.get("messages")
    template_key = str(data.get("template_key", "form") or "form").strip()
    model = str(data.get("model", "") or "").strip()

    if template_key not in TEMPLATE_CATALOG:
        return jsonify({"success": False, "error": "非法的模板类型"}), 400
    if not isinstance(messages, list) or not messages:
        return jsonify({"success": False, "error": "缺少 messages"}), 400
    if not Config.resolve_llm_key(model):
        return jsonify({"success": False, "error": "模型 API Key 未配置"}), 500

    try:
        normalized_messages = _normalize_messages(messages)
        payload = [{"role": "system", "content": _system_prompt_for(template_key)}]
        payload.extend(normalized_messages)

        completion = _request_completion(
            payload,
            max_tokens=Config.LLM_REQUEST_MAX_TOKENS,
            timeout_seconds=Config.LLM_REQUEST_TIMEOUT,
            model=model,
        )
        analysis = _finalize_analysis(completion["content"], template_key)
        return jsonify({
            "success": True,
            "analysis": analysis,
            "model": _resolve_model(model),
        })
    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)}), 500


@apex_ai_bp.route("/api/apex-ai/analyze-stream", methods=["POST"])
def analyze_stream():
    data = request.get_json(silent=True) or {}
    messages = data.get("messages")
    template_key = str(data.get("template_key", "form") or "form").strip()
    model = str(data.get("model", "") or "").strip()

    if template_key not in TEMPLATE_CATALOG:
        return jsonify({"success": False, "error": "非法的模板类型"}), 400
    if not isinstance(messages, list) or not messages:
        return jsonify({"success": False, "error": "缺少 messages"}), 400
    if not Config.resolve_llm_key(model):
        return jsonify({"success": False, "error": "模型 API Key 未配置"}), 500

    @stream_with_context
    def generate() -> Iterator[str]:
        try:
            normalized_messages = _normalize_messages(messages)
            payload = [{"role": "system", "content": _system_prompt_for(template_key)}]
            payload.extend(normalized_messages)

            yield _sse_event("status", {"message": "正在解析需求描述..."})

            response = _request_completion(
                payload,
                max_tokens=Config.LLM_REQUEST_MAX_TOKENS,
                timeout_seconds=Config.LLM_REQUEST_TIMEOUT,
                model=model,
                stream=True,
            )

            raw_content = ""
            chunk_count = 0

            for chunk in _iter_stream(response):
                delta = chunk.get("content", "")
                if not delta:
                    continue
                raw_content += delta
                chunk_count += 1
                if chunk_count == 1:
                    yield _sse_event("status", {"message": "模型开始返回内容..."})
                elif chunk_count % 12 == 0:
                    yield _sse_event("status", {"message": f"接收中，第 {chunk_count} 段"})

            analysis = _finalize_analysis(raw_content, template_key)
            yield _sse_event("result", {
                "success": True,
                "analysis": analysis,
                "model": _resolve_model(model),
            })
        except Exception as exc:
            yield _sse_event("error", {"error": str(exc)})
        finally:
            yield _sse_event("done", {"success": True})

    response = Response(generate(), content_type="text/event-stream; charset=utf-8")
    response.headers["Cache-Control"] = "no-cache"
    response.headers["X-Accel-Buffering"] = "no"
    return response


@apex_ai_bp.route("/api/apex-ai/generate", methods=["POST"])
def generate_doc():
    data = request.get_json(silent=True) or {}
    analysis = data.get("analysis")
    template_key = str(data.get("template_key", "") or "").strip()

    if not isinstance(analysis, dict) or not analysis:
        return jsonify({"success": False, "error": "缺少 analysis"}), 400
    if template_key not in TEMPLATE_CATALOG:
        template_key = str(analysis.get("template_key", "form") or "form").strip()
    if template_key not in TEMPLATE_CATALOG:
        return jsonify({"success": False, "error": "非法的模板类型"}), 400

    try:
        normalized = _normalize_analysis(analysis, template_key)
        docx_path = _build_docx(normalized, template_key)
        return jsonify({
            "success": True,
            "docx_name": docx_path.name,
            "title": normalized.get("document_title"),
            "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
            "download_url": f"/api/apex-ai/download/{docx_path.name}",
        })
    except Exception as exc:
        return jsonify({"success": False, "error": str(exc)}), 500


@apex_ai_bp.route("/api/apex-ai/exports", methods=["GET"])
def list_exports():
    EXPORTS_DIR.mkdir(parents=True, exist_ok=True)
    items = []
    for docx_path in sorted(EXPORTS_DIR.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True):
        meta_path = docx_path.with_suffix(".json")
        meta = {}
        if meta_path.exists():
            try:
                meta = json.loads(meta_path.read_text(encoding="utf-8"))
            except Exception:
                meta = {}
        items.append({
            "docx_name": docx_path.name,
            "title": meta.get("title") or docx_path.stem,
            "template_key": meta.get("template_key") or "",
            "template_label": meta.get("template_label") or "",
            "generated_at": meta.get("generated_at") or datetime.fromtimestamp(docx_path.stat().st_mtime).strftime("%Y-%m-%d %H:%M:%S"),
            "download_url": f"/api/apex-ai/download/{docx_path.name}",
        })
    return jsonify({"success": True, "items": items, "total": len(items)})


@apex_ai_bp.route("/api/apex-ai/download/<path:docx_name>", methods=["GET"])
def download(docx_name: str):
    safe_name = Path(docx_name).name
    docx_path = EXPORTS_DIR / safe_name
    if not docx_path.exists():
        return jsonify({"success": False, "error": "文件不存在"}), 404
    return send_file(
        docx_path,
        as_attachment=True,
        download_name=safe_name,
        mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

