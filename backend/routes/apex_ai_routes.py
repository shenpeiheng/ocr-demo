"""
Apex AI 需求文档生成接口 - 完全基于真实模板结构重构

基于 template_analysis_data_1.json 和 template_analysis_data_2.json 的分析结果
"""

from __future__ import annotations

import json
import logging
import re
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List
from uuid import uuid4

import requests
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docxtpl import DocxTemplate
from flask import Blueprint, jsonify, request, send_file

from config import Config

# 配置日志
logger = logging.getLogger(__name__)
logger.setLevel(logging.INFO)

# 创建日志目录
LOG_DIR = Path(__file__).resolve().parents[2] / "logs"
LOG_DIR.mkdir(parents=True, exist_ok=True)

# 日志格式
formatter = logging.Formatter(
    '%(asctime)s | %(levelname)s | %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

# 文件处理器 - 记录所有 AI 调用
ai_log_file = LOG_DIR / f"apex_ai_{datetime.now().strftime('%Y%m%d')}.log"
file_handler = logging.FileHandler(ai_log_file, encoding='utf-8')
file_handler.setLevel(logging.INFO)
file_handler.setFormatter(formatter)
logger.addHandler(file_handler)

# 控制台处理器 - 同时输出到控制台
console_handler = logging.StreamHandler()
console_handler.setLevel(logging.INFO)
console_handler.setFormatter(formatter)
logger.addHandler(console_handler)

# 阻止日志传播到父 logger（避免重复输出）
logger.propagate = False

apex_ai_bp = Blueprint("apex_ai", __name__)

ROOT_DIR = Path(__file__).resolve().parents[2]
APEX_AI_DIR = ROOT_DIR / "frontend" / "apex_ai"
EXPORTS_DIR = APEX_AI_DIR / "exports"
EXPORTS_DIR.mkdir(parents=True, exist_ok=True)

# 模板配置
TEMPLATES = {
    "form": {
        "key": "form",
        "label": "表单类（MD.050）",
        "file": "MD.050-JUSENERP-ERP-INV-库存出入库平台V2.0(1)_tpl.docx",  # 使用带占位符的模板
        "description": "适合业务表单类需求",
    },
    "report": {
        "key": "report",
        "label": "报表类（MD.050）",
        "file": "YX&HB-ERP-MD050-HBCUX.离散任务缺料报表(1)_tpl.docx",
        "description": "适合报表类需求",
    },
}


# ==================== 工具函数 ====================


def _resolve_model(model_alias: str = "") -> str:
    return Config.resolve_llm_model(model_alias)


def _normalize_chat_url(model_key: str = "") -> str:
    base_url = Config.resolve_llm_url(model_key)
    if base_url.endswith("/chat/completions"):
        return base_url
    return f"{base_url.rstrip('/')}/chat/completions"


def _call_llm(messages: List[Dict[str, str]], model: str = "", max_tokens: int = 8000) -> str:
    """调用 LLM 并返回文本内容"""
    api_key = Config.resolve_llm_key(model)
    resolved_model = _resolve_model(model)

    payload = {
        "model": resolved_model,
        "messages": messages,
        "temperature": 0.2,
        "max_tokens": max_tokens,
    }

    # 记录请求日志
    logger.info(f"LLM 请求 | 模型: {resolved_model} | 消息数: {len(messages)} | max_tokens: {max_tokens}")

    try:
        start_time = datetime.now()
        response = requests.post(
            _normalize_chat_url(model),
            headers={
                "Authorization": f"Bearer {api_key}",
                "Content-Type": "application/json",
            },
            json=payload,
            timeout=120,
        )
        elapsed = (datetime.now() - start_time).total_seconds()

        response.raise_for_status()
        data = response.json()
        choice = (data.get("choices") or [{}])[0]
        message = choice.get("message") or {}
        content = message.get("content", "") or ""

        # 记录响应日志
        usage = data.get("usage", {})
        logger.info(
            f"LLM 响应 | 耗时: {elapsed:.2f}s | "
            f"tokens: {usage.get('total_tokens', 0)} "
            f"(prompt: {usage.get('prompt_tokens', 0)}, completion: {usage.get('completion_tokens', 0)}) | "
            f"响应长度: {len(content)}"
        )

        return content
    except Exception as exc:
        logger.error(f"LLM 调用失败 | 模型: {resolved_model} | 错误: {exc}")
        raise RuntimeError(f"LLM 调用失败: {exc}") from exc


def _extract_json(text: str) -> Dict[str, Any]:
    """从 LLM 响应中提取 JSON"""
    text = (text or "").strip()
    if not text:
        return {}

    try:
        return json.loads(text)
    except json.JSONDecodeError:
        pass

    fenced_match = re.search(r"```(?:json)?\s*(\{.*?\})\s*```", text, re.S)
    if fenced_match:
        try:
            return json.loads(fenced_match.group(1))
        except json.JSONDecodeError:
            pass

    brace_index = text.find("{")
    if brace_index >= 0:
        from json import JSONDecoder
        decoder = JSONDecoder()
        try:
            parsed, _ = decoder.raw_decode(text[brace_index:])
            if isinstance(parsed, dict):
                return parsed
        except json.JSONDecodeError:
            pass

    return {}


# ==================== 提示词 ====================


def _build_form_prompt() -> str:
    """表单类文档的提示词"""
    return """你是 MD.050 表单类需求文档生成助手。

**任务**：用户输入简单需求，你主动推理并生成完整的文档内容。

**核心原则**：
- 基于常见 ERP 业务场景补充合理的字段、按钮、规则
- 不追问用户，直接生成完整初稿
- 所有字段都要填充完整内容，不要留空

**输出 JSON**：
{
  "doc_info": {
    "title": "文档标题",
    "project": "项目名",
    "module": "模块代码",
    "version": "版本号"
  },
  "overview": {
    "background": "业务背景（2-4段）",
    "requirements": ["需求1", "需求2", "需求3"],
    "problems": ["问题1", "问题2"],
    "usage": {
      "users": "使用对象",
      "responsibility": "所属职责",
      "menu": "菜单路径",
      "prerequisite": "前提条件"
    }
  },
  "design": {
    "process": "业务流程说明（2-4段）",
    "screen_summary": "界面列表概述",
    "screen_overview": "界面概述"
  },
  "screens": [
    {
      "name": "界面名称",
      "type": "维护|查询|审批",
      "purpose": "界面用途",
      "fields": [
        {"name": "字段名", "type": "类型", "required": "是|否", "desc": "说明"}
      ],
      "actions": ["按钮1：功能", "按钮2：功能"],
      "rules": ["规则1", "规则2"]
    }
  ],
  "issues": {
    "open": [
      {"problem": "未决问题", "solution": "", "owner": "负责人", "target_date": "目标日期", "actual_date": ""}
    ],
    "closed": [
      {"problem": "已决问题", "solution": "解决方案", "owner": "负责人", "target_date": "目标日期", "actual_date": "实际日期"}
    ]
  },
  "reply": "给用户的回复"
}

**要求**：
- 严格输出 JSON，不要 Markdown 代码块
- 界面通常 2-6 个
- 每个界面至少 5-10 个字段、3-5 个按钮、3-5 条规则
"""


def _build_report_prompt() -> str:
    """报表类文档的提示词"""
    return """你是 MD.050 报表类需求文档生成助手。

**任务**：用户输入简单需求，你主动推理并生成完整的报表文档内容。

**核心原则**：
- 基于常见 ERP 报表场景补充合理的参数、字段
- 不追问用户，直接生成完整初稿
- 所有字段都要填充完整内容

**输出 JSON**：
{
  "doc_info": {
    "title": "报表标题",
    "project": "项目名",
    "module": "模块代码",
    "version": "版本号"
  },
  "overview": {
    "background": "业务背景（2-4段）",
    "requirements": ["相关业务需求1"],
    "problems": ["开发后解决的问题1"],
    "usage": {"users": "使用对象", "responsibility": "所属职责", "menu": "菜单位置", "prerequisite": "前提条件"},
    "purpose": "报表目的",
    "scope": "适用范围",
    "references": ["参考文档1"]
  },
  "description": {
    "requirements": "业务需求（2-3段）",
    "features": ["特性1", "特性2"]
  },
    "tech_spec": {
        "reference_report": "参考报表",
        "request_group": "请求组",
        "data_scope": "数据范围",
        "grouping": "分组规则",
        "sorting": "排序规则",
        "data_logic": "取数逻辑与处理流程（2-4段）",
        "layout": "报表布局说明",
        "fields": [
      {"seq": "1", "name": "字段名", "type": "类型", "source": "来源", "desc": "说明"}
    ]
  },
  "parameters": [
    {"code": "参数编码", "name": "参数名", "type": "类型", "required": "是|否", "default": "默认值", "value_set": "值集/LOV逻辑", "desc": "说明"}
  ],
  "test_points": ["测试要点1", "测试要点2"],
  "notes": "其他注释",
  "reply": "给用户的回复"
}

**要求**：
- 严格输出 JSON，不要 Markdown 代码块
- 字段至少 10-20 个
- 参数至少 5-10 个
- 测试要点至少 3-5 条
"""


# ==================== 路由 ====================


@apex_ai_bp.route("/templates", methods=["GET"])
def get_templates():
    """获取模板列表"""
    return jsonify({"templates": list(TEMPLATES.values())})


@apex_ai_bp.route("/chat", methods=["POST"])
def chat():
    """聊天接口"""
    try:
        payload = request.get_json() or {}
        template_key = payload.get("template", "form")
        model = payload.get("model", "")
        messages = payload.get("messages", [])

        if not messages:
            return jsonify({"error": "消息列表不能为空"}), 400

        # 构建系统提示词
        system_prompt = _build_report_prompt() if template_key == "report" else _build_form_prompt()

        # 构建完整消息
        full_messages = [{"role": "system", "content": system_prompt}]
        for msg in messages[-10:]:
            role = msg.get("role", "user")
            content = str(msg.get("content", "")).strip()
            if role in ["user", "assistant"] and content:
                full_messages.append({"role": role, "content": content})

        # 调用 LLM
        response_text = _call_llm(full_messages, model)
        data = _extract_json(response_text)

        if not data:
            return jsonify({"error": "无法解析 LLM 返回的 JSON"}), 500

        reply = data.get("reply", "已生成文档结构，请查看右侧预览。")

        return jsonify({
            "success": True,
            "data": data,
            "reply": reply,
            "template": template_key
        })

    except Exception as e:
        return jsonify({"error": str(e)}), 500


@apex_ai_bp.route("/preview", methods=["POST"])
def preview_document():
    """生成文档预览 HTML"""
    try:
        payload = request.get_json() or {}
        data = payload.get("data", {})
        template_key = payload.get("template", "form")

        html = _generate_report_preview(data) if template_key == "report" else _generate_form_preview(data)

        return jsonify({"success": True, "html": html, "data": data, "template_key": template_key})

    except Exception as e:
        logger.error(f"生成预览失败: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


@apex_ai_bp.route("/generate", methods=["POST"])
def generate_document():
    """生成 Word 文档"""
    try:
        payload = request.get_json() or {}
        data = payload.get("analysis", {})  # 前端发送的是 analysis，不是 data
        template_key = payload.get("template_key", "form")  # 前端发送的是 template_key

        # 加载模板
        template_file = APEX_AI_DIR / TEMPLATES[template_key]["file"]
        if not template_file.exists():
            return jsonify({"success": False, "error": f"模板文件不存在: {template_file}"}), 404

        # 使用 docxtpl 加载模板
        tpl = DocxTemplate(str(template_file))

        # 准备渲染数据
        doc_info = data.get("doc_info", {})
        current_date = datetime.now().strftime("%Y年%m月%d日")

        overview = data.get("overview", {})
        design = data.get("design", {})
        description = data.get("description", {})
        tech_spec = data.get("tech_spec", {})
        parameters = data.get("parameters", [])
        current_month = datetime.now().strftime("%Y年%m月")

        def lines(value: Any) -> str:
            if isinstance(value, list):
                return "\n".join(f"{index}. {item}" for index, item in enumerate(value, 1))
            return str(value or "")

        def normalize_issues(values: Any) -> List[Dict[str, str]]:
            normalized = []
            for value in values or []:
                if isinstance(value, dict):
                    normalized.append({
                        "problem": str(value.get("problem") or value.get("issue") or ""),
                        "solution": str(value.get("solution") or ""),
                        "owner": str(value.get("owner") or "待确认"),
                        "target_date": str(value.get("target_date") or "待确认"),
                        "actual_date": str(value.get("actual_date") or ""),
                    })
                else:
                    normalized.append({
                        "problem": str(value), "solution": "", "owner": "待确认",
                        "target_date": "待确认", "actual_date": "",
                    })
            return normalized

        issues = data.get("issues", {})

        context = {
            "project": doc_info.get("project", ""),
            "module": doc_info.get("module", ""),
            "project_type": "ERP实施项目",
            "title": doc_info.get("title", ""),
            "document_type": "功能设计文档",
            "change_date": current_date,
            "created_month": current_month,
            "author": "SIE顾问",
            "author_role": "业务顾问",
            "approver": doc_info.get("approver", "待确认"),
            "project_manager": doc_info.get("project_manager", "待确认"),
            "sie_project_manager": doc_info.get("sie_project_manager", "待确认"),
            "reviewer_name": doc_info.get("reviewer_name", "待确认"),
            "reviewer_role": doc_info.get("reviewer_role", "项目负责人"),
            "reviewer_signature": doc_info.get("reviewer_signature", ""),
            "distribution_no": doc_info.get("distribution_no", "1"),
            "distribution_name": doc_info.get("distribution_name", doc_info.get("project", "项目组")),
            "distribution_location": doc_info.get("distribution_location", "项目文档库"),
            "version": doc_info.get("version", "1.0"),
            "change_ref": "初稿",
            "preface": f"本文档为{doc_info.get('project', '项目')}的{doc_info.get('title', '功能')}功能设计说明书，详细描述业务需求、界面设计、字段定义及业务规则。",
            "preface_revision_note": "本文档将根据业务确认、测试结果与用户反馈持续修订和完善。",
            "preface_development_note": "开发人员可依据本文档开展设计与开发，并与业务顾问及关键用户及时确认需求和校验结果。",
            "reference_document": "无",
            "key_users": "待确认",
            "consultants": "SIE顾问",
            "overview_background": lines(overview.get("background")),
            "overview_requirements": lines(overview.get("requirements")),
            "overview_problems": lines(overview.get("problems")),
            "report_background": lines(overview.get("background")),
            "report_problems": lines(overview.get("problems")),
            "report_usage": "；".join([
                f"使用对象：{overview.get('usage', {}).get('users', '暂无')}",
                f"所属职责：{overview.get('usage', {}).get('responsibility', '暂无')}",
                f"菜单位置：{overview.get('usage', {}).get('menu', '暂无')}",
                f"前提条件：{overview.get('usage', {}).get('prerequisite', '暂无')}",
            ]),
            "report_purpose": lines(overview.get("purpose")),
            "report_scope": lines(overview.get("scope")),
            "report_references": lines(overview.get("references")),
            "description_requirements": lines(description.get("requirements")),
            "description_features": lines(description.get("features")),
            "tech_reference_report": lines(tech_spec.get("reference_report")),
            "tech_request_group": lines(tech_spec.get("request_group")),
            "tech_data_scope": lines(tech_spec.get("data_scope")),
            "tech_grouping": lines(tech_spec.get("grouping")),
            "tech_sorting": lines(tech_spec.get("sorting")),
            "tech_data_logic": lines(tech_spec.get("data_logic")),
            "report_layout": lines(tech_spec.get("layout")),
            "report_fields": tech_spec.get("fields", []),
            "parameters": parameters,
            "test_points": lines(data.get("test_points")),
            "notes": lines(data.get("notes")),
            "parameter_notes": lines(data.get("parameter_notes")),
            "design_process": lines(design.get("process")),
            "design_overview": lines(design.get("screen_overview") or design.get("screen_summary")),
            "screens": data.get("screens", []),
            "open_issues": normalize_issues(issues.get("open")),
            "closed_issues": normalize_issues(issues.get("closed")),
            "usage_users": lines(overview.get("usage", {}).get("users")),
            "usage_responsibility": lines(overview.get("usage", {}).get("responsibility")),
            "usage_menu": lines(overview.get("usage", {}).get("menu")),
            "usage_prerequisite": lines(overview.get("usage", {}).get("prerequisite")),
        }

        # 渲染模板
        tpl.render(context)

        # 获取渲染后的文档对象继续填充其他内容
        doc = tpl.docx

        # 生成文档
        if template_key == "report":
            _fill_report_template(doc, data)
            _remove_report_static_toc(doc)
        else:
            _fill_form_template(doc, data)

        if template_key == "form":
            _remove_toc(doc)

        # 保留模板中的 Word TOC 域，但不在打开时触发域更新。Word 对 updateFields
        # 和 dirty TOC 都可能弹出外部域安全确认；准确目录应由排版引擎在生成阶段刷新。
        _disable_update_fields(doc)

        # 保存文档
        title = doc_info.get('title', '需求文档')
        filename = f"{title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        output_path = EXPORTS_DIR / filename

        tpl.save(str(output_path))

        # 保存预览数据（JSON）
        preview_data = {
            "template_key": template_key,
            "data": data,
            "generated_at": datetime.now().isoformat()
        }
        json_filename = filename.replace('.docx', '.json')
        json_path = EXPORTS_DIR / json_filename
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(preview_data, f, ensure_ascii=False, indent=2)

        return jsonify({
            "success": True,
            "title": title,
            "filename": filename,
            "download_url": f"/apex_ai/exports/{filename}"
        })

    except Exception as e:
        logger.error(f"生成文档失败: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


# ==================== HTML 预览生成 ====================


def _generate_form_preview(data: Dict[str, Any]) -> str:
    """生成表单类文档的 HTML 预览"""
    doc_info = data.get("doc_info", {})
    overview = data.get("overview", {})
    design = data.get("design", {})
    screens = data.get("screens", [])
    issues = data.get("issues", {})

    parts = [
        '<div style="font-family: 微软雅黑; padding: 20px; background: #fff; line-height: 1.8;">',

        # 文档头部
        '<div style="text-align: center; margin-bottom: 30px; border-bottom: 2px solid #ad063e; padding-bottom: 20px;">',
        f'<h1 style="font-size: 24px; color: #303133; margin: 10px 0;">{doc_info.get("title", "")}</h1>',
        f'<p style="color: #909399; font-size: 14px;">{doc_info.get("project", "")} | {doc_info.get("module", "")} | {doc_info.get("version", "")}</p>',
        '</div>',

        # 3. 概述
        '<h2 style="font-size: 18px; color: #fff; background: #ad063e; padding: 8px 12px; margin: 30px 0 15px;">3. 概述</h2>',

        '<h3 style="font-size: 14px; color: #606266; margin: 15px 0 8px;">3.1 业务背景</h3>',
        f'<div style="text-indent: 2em; margin-bottom: 15px;">{overview.get("background", "")}</div>',

        '<h3 style="font-size: 14px; color: #606266; margin: 15px 0 8px;">3.2 相关业务需求</h3>',
        '<ol style="margin-left: 20px;">',
    ]

    for req in overview.get("requirements", []):
        parts.append(f'<li>{req}</li>')
    parts.append('</ol>')

    parts.append('<h3 style="font-size: 14px; color: #606266; margin: 15px 0 8px;">3.3 开发后解决的问题</h3>')
    parts.append('<ol style="margin-left: 20px;">')
    for prob in overview.get("problems", []):
        parts.append(f'<li>{prob}</li>')
    parts.append('</ol>')

    usage = overview.get("usage") or {}
    parts.append('<h3 style="font-size: 14px; color: #606266; margin: 15px 0 8px;">3.4 使用说明</h3>')
    parts.append(f'<p><strong>使用对象：</strong>{usage.get("users", "")}</p>')
    parts.append(f'<p><strong>所属职责：</strong>{usage.get("responsibility", "")}</p>')
    parts.append(f'<p><strong>菜单路径：</strong>{usage.get("menu", "")}</p>')
    parts.append(f'<p><strong>前提条件：</strong>{usage.get("prerequisite", "")}</p>')

    # 4. 总体结构设计
    parts.append('<h2 style="font-size: 18px; color: #fff; background: #ad063e; padding: 8px 12px; margin: 30px 0 15px;">4. 总体结构设计</h2>')
    parts.append('<h3 style="font-size: 14px; color: #606266; margin: 15px 0 8px;">4.1 相关业务流程说明</h3>')
    parts.append(f'<div style="text-indent: 2em; margin-bottom: 15px;">{design.get("process", "")}</div>')

    # 5. 详细功能设计
    parts.append('<h2 style="font-size: 18px; color: #fff; background: #ad063e; padding: 8px 12px; margin: 30px 0 15px;">5. 详细功能设计</h2>')

    for i, screen in enumerate(screens, 1):
        parts.append(f'<h3 style="font-size: 14px; color: #606266; margin: 20px 0 10px;">界面{i} - {screen.get("name", "")}</h3>')
        parts.append(f'<p><strong>类型：</strong>{screen.get("type", "")} | <strong>用途：</strong>{screen.get("purpose", "")}</p>')

        parts.append('<p style="margin-top: 15px;"><strong>字段清单：</strong></p>')
        parts.append('<table style="width: 100%; border-collapse: collapse; margin-bottom: 15px;">')
        parts.append('<tr style="background: #f5f7fa;"><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">字段名</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">类型</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">必填</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">说明</th></tr>')
        for field in screen.get("fields", []):
            parts.append(f'<tr><td style="border: 1px solid #ddd; padding: 8px;">{field.get("name", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{field.get("type", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{field.get("required", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{field.get("desc", "")}</td></tr>')
        parts.append('</table>')

        parts.append('<p><strong>按钮功能：</strong></p>')
        parts.append('<ul style="margin-left: 20px;">')
        for action in screen.get("actions", []):
            parts.append(f'<li>{action}</li>')
        parts.append('</ul>')

        parts.append('<p><strong>业务规则：</strong></p>')
        parts.append('<ul style="margin-left: 20px;">')
        for rule in screen.get("rules", []):
            parts.append(f'<li>{rule}</li>')
        parts.append('</ul>')

    # 6. 未决与已决问题
    parts.append('<h2 style="font-size: 18px; color: #fff; background: #ad063e; padding: 8px 12px; margin: 30px 0 15px;">6. 未决与已决问题</h2>')

    parts.append('<h3 style="font-size: 14px; color: #606266; margin: 15px 0 8px;">6.1 未决问题</h3>')
    open_issues = issues.get("open", [])
    if open_issues:
        parts.append('<ol style="margin-left: 20px;">')
        for issue in open_issues:
            parts.append(f'<li>{issue}</li>')
        parts.append('</ol>')
    else:
        parts.append('<p>暂无未决问题。</p>')

    parts.append('<h3 style="font-size: 14px; color: #606266; margin: 15px 0 8px;">6.2 已决问题</h3>')
    closed_issues = issues.get("closed", [])
    if closed_issues:
        parts.append('<ol style="margin-left: 20px;">')
        for issue in closed_issues:
            parts.append(f'<li>{issue}</li>')
        parts.append('</ol>')
    else:
        parts.append('<p>暂无已决问题。</p>')

    parts.append('</div>')
    return ''.join(parts)


def _generate_report_preview(data: Dict[str, Any]) -> str:
    """生成报表类文档的 HTML 预览"""
    doc_info = data.get("doc_info", {})
    overview = data.get("overview", {})
    description = data.get("description", {})
    tech_spec = data.get("tech_spec", {})
    parameters = data.get("parameters", [])
    test_points = data.get("test_points", [])

    parts = [
        '<div style="font-family: 微软雅黑; padding: 20px; background: #fff; line-height: 1.8;">',

        '<div style="text-align: center; margin-bottom: 30px; border-bottom: 2px solid #ad063e; padding-bottom: 20px;">',
        f'<h1 style="font-size: 24px; color: #303133; margin: 10px 0;">{doc_info.get("title", "")}</h1>',
        f'<p style="color: #909399; font-size: 14px;">{doc_info.get("project", "")} | {doc_info.get("module", "")} | {doc_info.get("version", "")}</p>',
        '</div>',

        '<h2 style="font-size: 18px; color: #fff; background: #ad063e; padding: 8px 12px; margin: 30px 0 15px;">2. 概述</h2>',
        f'<p><strong>目的：</strong>{overview.get("purpose", "")}</p>',
        f'<p><strong>适用范围：</strong>{overview.get("scope", "")}</p>',

        '<h2 style="font-size: 18px; color: #fff; background: #ad063e; padding: 8px 12px; margin: 30px 0 15px;">3. 详细描述</h2>',
        f'<div style="text-indent: 2em; margin-bottom: 15px;">{description.get("requirements", "")}</div>',

        '<p><strong>主要特性：</strong></p>',
        '<ul style="margin-left: 20px;">',
    ]

    for feature in description.get("features", []):
        parts.append(f'<li>{feature}</li>')
    parts.append('</ul>')

    parts.append('<h2 style="font-size: 18px; color: #fff; background: #ad063e; padding: 8px 12px; margin: 30px 0 15px;">4. 技术规格</h2>')
    parts.append(f'<p><strong>参考报表：</strong>{tech_spec.get("reference_report", "")}</p>')
    parts.append(f'<p><strong>数据范围：</strong>{tech_spec.get("data_scope", "")}</p>')
    parts.append(f'<p><strong>分组规则：</strong>{tech_spec.get("grouping", "")}</p>')
    parts.append(f'<p><strong>排序规则：</strong>{tech_spec.get("sorting", "")}</p>')

    parts.append('<p style="margin-top: 15px;"><strong>字段说明：</strong></p>')
    parts.append('<table style="width: 100%; border-collapse: collapse; margin-bottom: 20px;">')
    parts.append('<tr style="background: #f5f7fa;"><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">序号</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">字段名</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">类型</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">来源</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">说明</th></tr>')
    for field in tech_spec.get("fields", []):
        parts.append(f'<tr><td style="border: 1px solid #ddd; padding: 8px;">{field.get("seq", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{field.get("name", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{field.get("type", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{field.get("source", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{field.get("desc", "")}</td></tr>')
    parts.append('</table>')

    parts.append('<h2 style="font-size: 18px; color: #fff; background: #ad063e; padding: 8px 12px; margin: 30px 0 15px;">6. 报表参数</h2>')
    parts.append('<table style="width: 100%; border-collapse: collapse; margin-bottom: 20px;">')
    parts.append('<tr style="background: #f5f7fa;"><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">参数名</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">类型</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">必填</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">默认值</th><th style="border: 1px solid #ddd; padding: 8px; text-align: left;">说明</th></tr>')
    for param in parameters:
        parts.append(f'<tr><td style="border: 1px solid #ddd; padding: 8px;">{param.get("name", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{param.get("type", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{param.get("required", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{param.get("default", "")}</td><td style="border: 1px solid #ddd; padding: 8px;">{param.get("desc", "")}</td></tr>')
    parts.append('</table>')

    parts.append('<h2 style="font-size: 18px; color: #fff; background: #ad063e; padding: 8px 12px; margin: 30px 0 15px;">8. 测试要点</h2>')
    parts.append('<ol style="margin-left: 20px;">')
    for point in test_points:
        parts.append(f'<li>{point}</li>')
    parts.append('</ol>')

    parts.append('</div>')
    return ''.join(parts)


# ==================== Word 文档填充 ====================


def _fill_form_template(doc: Document, data: Dict[str, Any]) -> None:
    """填充表单类模板 - 完整版"""
    doc_info = data.get("doc_info", {})
    overview = data.get("overview", {})
    design = data.get("design", {})
    screens = data.get("screens", [])
    issues = data.get("issues", {})
    usage = overview.get("usage", {})

    # 填充文档头部信息
    _safe_set_paragraph(doc, 1, doc_info.get("project", ""))
    _safe_set_paragraph(doc, 2, "ERP实施项目")
    _safe_set_paragraph(doc, 3, doc_info.get("title", ""))
    _safe_set_paragraph(doc, 4, "功能设计文档")
    _safe_set_paragraph(doc, 6, "作者:\tSIE顾问")
    _safe_set_paragraph(doc, 7, f"创建日期:\t{datetime.now().strftime('%Y年%m月')}")
    _safe_set_paragraph(doc, 8, f"最后更新日期：\t{datetime.now().strftime('%Y年%m月%d日')}")
    _safe_set_paragraph(doc, 9, f"版本：\t{doc_info.get('version', '1.0')}")

    # 填充前言章节（说明和参与人员）
    _fill_preface_section(doc, doc_info)

    # 填充界面设计 - 动态填充所有界面内容
    _fill_screens_content(doc, screens, overview, design, issues)

    logger.info(f"表单类模板填充完成：{len(screens)} 个界面")


def _fill_report_template(doc: Document, data: Dict[str, Any]) -> None:
    """报表类内容已由 docxtpl 直接写入模板，避免在文末追加章节。"""
    logger.info("报表类模板渲染完成")


def _safe_set_paragraph(doc: Document, index: int, text: str) -> None:
    """安全地设置段落文本，保持原有样式"""
    try:
        if index < len(doc.paragraphs):
            # 保存原有样式
            para = doc.paragraphs[index]
            # 清空现有内容
            para.clear()
            # 添加新内容（保持原有字体格式）
            run = para.add_run(text)
            # 可以在这里设置字体，但为了保持模板样式，暂时不设置
        else:
            logger.warning(f"段落索引 {index} 超出范围，跳过")
    except Exception as e:
        logger.error(f"设置段落 {index} 失败: {e}")


def _fill_preface_section(doc: Document, doc_info: Dict) -> None:
    """填充前言章节：2.1 说明 和 2.2 参与人员"""
    try:
        current_date = datetime.now().strftime('%Y年%m月%d日')
        project_name = doc_info.get('project', '项目')
        title = doc_info.get('title', '功能')

        # 查找"2.1 说明"或"说明"段落
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()

            # 查找2.1说明章节，并在其后填充说明内容
            if "2.1" in para_text and "说明" in para_text:
                # 在下一段填充说明内容
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    # 如果下一段是空的或者是模板占位文本，则替换
                    if not next_para.text.strip() or "本文档" in next_para.text:
                        next_para.clear()
                        next_para.add_run(f"本文档为{project_name}的{title}功能设计说明书，详细描述了业务需求、界面设计、字段定义及业务规则等内容。")
                        logger.info("已填充2.1说明内容")
                break

        # 查找"2.2 参与人员"段落
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()

            # 查找2.2参与人员章节
            if "2.2" in para_text and "参与人员" in para_text:
                # 查找该章节后的表格
                # 通常参与人员是一个表格，我们需要填充日期列
                found_table = False
                for table in doc.tables:
                    # 检查表格是否在该段落之后且包含"姓名"、"角色"等字段
                    if len(table.rows) > 0:
                        header_text = " ".join([cell.text for cell in table.rows[0].cells])
                        if "姓名" in header_text or "角色" in header_text or "职责" in header_text:
                            # 找到日期列
                            date_col_index = -1
                            for col_idx, cell in enumerate(table.rows[0].cells):
                                if "日期" in cell.text:
                                    date_col_index = col_idx
                                    break

                            # 填充所有数据行的日期
                            if date_col_index >= 0:
                                for row_idx in range(1, len(table.rows)):
                                    if date_col_index < len(table.rows[row_idx].cells):
                                        table.rows[row_idx].cells[date_col_index].text = current_date

                                found_table = True
                                logger.info(f"已填充2.2参与人员表格日期: {current_date}")
                                break

                if found_table:
                    break

    except Exception as e:
        logger.error(f"填充前言章节失败: {e}", exc_info=True)


def _fill_usage_section(doc: Document, usage: Dict) -> None:
    """填充3.4使用说明章节"""
    try:
        # 查找"3.4 使用说明"或"使用说明"段落
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()

            # 查找3.4使用说明章节
            if "3.4" in para_text and "使用说明" in para_text:
                # 查找该章节后的表格
                found_table = False
                for table_idx, table in enumerate(doc.tables):
                    if len(table.rows) > 0:
                        # 检查表格是否包含"使用对象"、"所属职责"等字段
                        table_text = " ".join([cell.text for cell in table.rows[0].cells if cell.text])

                        # 查找包含使用说明相关字段的表格
                        if "使用对象" in table_text or "所属职责" in table_text or "菜单" in table_text:
                            # 填充使用说明表格
                            for row_idx in range(len(table.rows)):
                                for cell in table.rows[row_idx].cells:
                                    cell_text = cell.text.strip()

                                    # 根据单元格内容判断要填充的值
                                    if "使用对象" in cell_text and ":" in cell_text:
                                        cell.text = f"使用对象: {usage.get('users', '')}"
                                    elif "所属职责" in cell_text and ":" in cell_text:
                                        cell.text = f"所属职责: {usage.get('responsibility', '')}"
                                    elif "菜单路径" in cell_text and ":" in cell_text:
                                        cell.text = f"菜单路径: {usage.get('menu', '')}"
                                    elif "前提条件" in cell_text and ":" in cell_text:
                                        cell.text = f"前提条件: {usage.get('prerequisite', '')}"

                            found_table = True
                            logger.info("已填充3.4使用说明表格")
                            break

                if found_table:
                    break

    except Exception as e:
        logger.error(f"填充使用说明章节失败: {e}", exc_info=True)


def _fill_design_sections(doc: Document, design: Dict, screens: List[Dict]) -> None:
    """填充4.2界面列表和4.3流程设计"""
    try:
        # 填充4.2界面列表
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()

            # 查找4.2界面列表章节
            if "4.2" in para_text and ("界面列表" in para_text or "界面" in para_text):
                # 在下一段填充界面列表内容
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    next_para.clear()

                    # 生成界面列表文本
                    screen_list = []
                    for idx, screen in enumerate(screens, 1):
                        screen_name = screen.get('name', '')
                        screen_type = screen.get('type', '')
                        screen_purpose = screen.get('purpose', '')
                        screen_list.append(f"{idx}. {screen_name}（{screen_type}）：{screen_purpose}")

                    if screen_list:
                        next_para.add_run("\n".join(screen_list))
                        logger.info(f"已填充4.2界面列表：{len(screens)}个界面")
                    else:
                        next_para.add_run("根据业务需求，本功能包含若干界面，详见第5章详细功能设计。")
                break

        # 填充4.3流程设计
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()

            # 查找4.3或流程设计章节
            if ("4.3" in para_text or "4.1" in para_text) and ("流程" in para_text or "业务流程" in para_text):
                # 在下一段填充流程设计内容
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    # 如果下一段是空的或者是模板占位文本，则替换
                    if not next_para.text.strip() or len(next_para.text.strip()) < 20:
                        next_para.clear()
                        process_text = design.get('process', '')
                        if process_text:
                            next_para.add_run(process_text)
                            logger.info("已填充4.3流程设计内容")
                        else:
                            next_para.add_run("业务流程详见各界面操作说明。")
                break

    except Exception as e:
        logger.error(f"填充设计章节失败: {e}", exc_info=True)


def _fill_screens_content(doc: Document, screens: List[Dict], overview: Dict, design: Dict, issues: Dict) -> None:
    """填充界面内容到 Word 文档 - 按照标准MD.050格式"""
    try:
        # Word 的自动编号不包含在 paragraph.text 中，因此按章节名称定位。
        section_5_index = -1
        for i, para in enumerate(doc.paragraphs):
            if "详细功能设计" in para.text:
                section_5_index = i
                break

        if section_5_index == -1:
            logger.warning("未找到'5. 详细功能设计'章节")
            return

        # 删除第5章后面的所有旧内容（保留章节标题）
        # 从第5章标题的下一段开始删除，直到找到第6章
        section_6_index = -1
        for i in range(section_5_index + 1, len(doc.paragraphs)):
            if "问题" in doc.paragraphs[i].text and (
                "未结" in doc.paragraphs[i].text or "未决" in doc.paragraphs[i].text
            ):
                section_6_index = i
                break

        if section_6_index == -1:
            logger.warning("未找到问题章节，无法定位详细设计插入位置")
            return

        section_5 = doc.paragraphs[section_5_index]._element
        section_6 = doc.paragraphs[section_6_index]._element

        # 清除示例模板第5章内容，但保留第5、6章标题及第6章的问题表。
        element = section_5.getnext()
        while element is not None and element is not section_6:
            next_element = element.getnext()
            element.getparent().remove(element)
            element = next_element

        def add_paragraph(*args, **kwargs):
            paragraph = doc.add_paragraph(*args, **kwargs)
            section_6.addprevious(paragraph._element)
            return paragraph

        def add_table(*args, **kwargs):
            table = doc.add_table(*args, **kwargs)
            section_6.addprevious(table._element)
            return table

        def set_manual_heading(paragraph, style: str, text: str):
            """使用手工章节号，同时保留标题级别供目录识别。"""
            paragraph.style = style
            p_pr = paragraph._p.get_or_add_pPr()
            num_pr = p_pr.find(qn("w:numPr"))
            if num_pr is None:
                num_pr = OxmlElement("w:numPr")
                p_pr.append(num_pr)
            num_id = num_pr.find(qn("w:numId"))
            if num_id is None:
                num_id = OxmlElement("w:numId")
                num_pr.append(num_id)
            num_id.set(qn("w:val"), "0")
            return paragraph.add_run(text)

        # 为每个界面生成详细设计
        for idx, screen in enumerate(screens, 1):
            # 每个界面只占一个二级章节；编号由 Word 标题样式统一处理。
            p = add_paragraph()
            run = set_manual_heading(p, 'Heading 2', f"5.{idx} {screen.get('name', '')}")
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0, 0, 0)

            # 三级标题按 5.x.1、5.x.2 展开，并进入自动目录。
            p = add_paragraph()
            run = set_manual_heading(p, 'Heading 3', f"5.{idx}.1 界面设计")
            run.font.size = Pt(12)
            run.font.bold = True

            p = add_paragraph()
            p.add_run(f"类型：{screen.get('type', '')}\n")
            p.add_run(f"用途：{screen.get('purpose', '')}")

            # 字段清单
            p = add_paragraph()
            run = set_manual_heading(p, 'Heading 3', f"5.{idx}.2 字段清单")
            run.font.size = Pt(12)
            run.font.bold = True

            # 创建字段表格
            fields = screen.get('fields', [])
            if fields:
                table = add_table(rows=1 + len(fields), cols=4)
                table.style = 'Table Grid'

                # 表头
                header_cells = table.rows[0].cells
                header_cells[0].text = '字段名'
                header_cells[1].text = '类型'
                header_cells[2].text = '必填'
                header_cells[3].text = '说明'

                # 表头加粗
                for cell in header_cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.bold = True

                # 填充字段数据
                for field_idx, field in enumerate(fields, 1):
                    row_cells = table.rows[field_idx].cells
                    row_cells[0].text = field.get('name', '')
                    row_cells[1].text = field.get('type', '')
                    row_cells[2].text = field.get('required', '')
                    row_cells[3].text = field.get('desc', '')
            else:
                p = add_paragraph()
                p.add_run("（暂无字段定义）")

            # 5.X.3 按钮功能
            actions = screen.get('actions', [])
            p = add_paragraph()
            run = set_manual_heading(p, 'Heading 3', f"5.{idx}.3 按钮功能")
            run.font.size = Pt(12)
            run.font.bold = True
            if actions:
                for action in actions:
                    p = add_paragraph()
                    p.add_run(f"• {action}")
            else:
                add_paragraph().add_run("（暂无按钮定义）")

            # 5.X.4 业务规则
            rules = screen.get('rules', [])
            p = add_paragraph()
            run = set_manual_heading(p, 'Heading 3', f"5.{idx}.4 业务规则")
            run.font.size = Pt(12)
            run.font.bold = True
            if rules:
                for rule in rules:
                    p = add_paragraph()
                    p.add_run(f"• {rule}")
            else:
                add_paragraph().add_run("（暂无业务规则）")

            # 界面原型（预留截图位置）
            p = add_paragraph()
            run = set_manual_heading(p, 'Heading 3', f"5.{idx}.5 界面原型")
            run.font.size = Pt(12)
            run.font.bold = True

            p = add_paragraph()
            p.add_run("（界面截图待补充）")

            # 添加分隔空行
            add_paragraph()

        logger.info(f"成功填充 {len(screens)} 个界面的详细设计")

    except Exception as e:
        logger.error(f"填充界面内容失败: {e}", exc_info=True)


def _fill_issues_section(doc: Document, issues: Dict) -> None:
    """填充6. 未决与已决问题章节"""
    try:
        open_issues = issues.get('open', [])
        closed_issues = issues.get('closed', [])

        # 查找6.1未决问题章节
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()

            if "问题" in para_text and ("未决" in para_text or "未结" in para_text) and "与" not in para_text:
                # 在下一段填充未决问题
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    next_para.clear()

                    if open_issues:
                        issue_list = []
                        for idx, issue in enumerate(open_issues, 1):
                            issue_list.append(f"{idx}. {issue}")
                        next_para.add_run("\n".join(issue_list))
                        logger.info(f"已填充6.1未决问题：{len(open_issues)}条")
                    else:
                        next_para.add_run("暂无未决问题。")
                break

        # 查找6.2已决问题章节
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()

            if "问题" in para_text and ("已决" in para_text or "已结" in para_text) and "与" not in para_text:
                # 在下一段填充已决问题
                if i + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[i + 1]
                    next_para.clear()

                    if closed_issues:
                        issue_list = []
                        for idx, issue in enumerate(closed_issues, 1):
                            issue_list.append(f"{idx}. {issue}")
                        next_para.add_run("\n".join(issue_list))
                        logger.info(f"已填充6.2已决问题：{len(closed_issues)}条")
                    else:
                        next_para.add_run("暂无已决问题。")
                break

    except Exception as e:
        logger.error(f"填充问题章节失败: {e}", exc_info=True)


def _add_toc_update_note(doc: Document) -> None:
    """在目录区域添加更新提示"""
    try:
        # 查找"目录"或"Table of Contents"标题
        toc_index = -1
        for i, para in enumerate(doc.paragraphs):
            para_text = para.text.strip()
            if "目录" in para_text or "Table of Contents" in para_text.lower():
                toc_index = i
                break

        if toc_index >= 0 and toc_index + 1 < len(doc.paragraphs):
            # 在目录标题后插入更新提示
            note_para = doc.paragraphs[toc_index + 1]
            note_para.clear()
            run = note_para.add_run('【提示】文档已生成，请在 Word 中右键点击目录，选择"更新域"→"更新整个目录"以刷新章节编号和页码。')
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(255, 0, 0)
            run.font.bold = True
            logger.info("已添加目录更新提示")

    except Exception as e:
        logger.error(f"添加目录更新提示失败: {e}", exc_info=True)


def _enable_update_fields(doc: Document) -> None:
    """让 Word 打开文档时自动更新目录、页码和其它域。"""
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")


def _disable_update_fields(doc: Document) -> None:
    """关闭打开文档时更新所有域，避免 Word 的外部域安全提示。"""
    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is not None:
        settings.remove(update_fields)


def _replace_with_static_toc(doc: Document, screens: List[Dict[str, Any]]) -> None:
    """删除 Word TOC 域，生成可跳转、无页码的动态目录。"""
    body = doc.element.body

    # 模板目录位于内容控件中；整个移除可一并清掉 TOC/PAGEREF 域和旧结果。
    for sdt in list(body.iter(qn("w:sdt"))):
        instructions = "".join(
            element.text or "" for element in sdt.iter(qn("w:instrText"))
        )
        if "TOC" in instructions.upper():
            parent = sdt.getparent()
            if parent is not None:
                parent.remove(sdt)

    preface = next((p for p in doc.paragraphs if p.text.strip() == "前言"), None)
    if preface is None:
        logger.warning("未找到前言章节，无法插入静态目录")
        return

    heading_targets = {
        "文档控制": "toc_1", "更改记录": "toc_1_2", "审核": "toc_1_3",
        "分发": "toc_1_4", "前言": "toc_2", "说明": "toc_2_1",
        "参与人员": "toc_2_2", "概述": "toc_3", "业务背景": "toc_3_1",
        "相关业务需求": "toc_3_2", "开发后解决的问题": "toc_3_3",
        "使用说明": "toc_3_4", "总体结构设计": "toc_4",
        "相关业务流程说明": "toc_4_1", "界面列表": "toc_4_2",
        "流程设计": "toc_4_3", "详细功能设计": "toc_5",
        "未结与已结问题": "toc_6", "未结问题": "toc_6_1", "已结问题": "toc_6_2",
    }
    bookmark_id = 100
    for paragraph in doc.paragraphs:
        target = heading_targets.get(paragraph.text.strip())
        if target:
            _add_bookmark(paragraph, target, bookmark_id)
            bookmark_id += 1

    entries = [
        (0, "目录", None),
        (0, "1 文档控制", "toc_1"),
        (1, "1.1 批准人", "toc_1"),
        (1, "1.2 更改记录", "toc_1_2"),
        (1, "1.3 审核", "toc_1_3"),
        (1, "1.4 分发", "toc_1_4"),
        (0, "2 前言", "toc_2"),
        (1, "2.1 说明", "toc_2_1"),
        (1, "2.2 参与人员", "toc_2_2"),
        (0, "3 概述", "toc_3"),
        (1, "3.1 业务背景", "toc_3_1"),
        (1, "3.2 相关业务需求", "toc_3_2"),
        (1, "3.3 开发后解决的问题", "toc_3_3"),
        (1, "3.4 使用说明", "toc_3_4"),
        (0, "4 总体结构设计", "toc_4"),
        (1, "4.1 相关业务流程说明", "toc_4_1"),
        (1, "4.2 界面列表", "toc_4_2"),
        (1, "4.3 流程设计", "toc_4_3"),
        (0, "5 详细功能设计", "toc_5"),
    ]
    for index, screen in enumerate(screens, 1):
        screen_name = screen.get("name", "")
        screen_target = f"toc_5_{index}"
        for paragraph in doc.paragraphs:
            if paragraph.text.strip() == f"5.{index} {screen_name}":
                _add_bookmark(paragraph, screen_target, bookmark_id)
                bookmark_id += 1
                break
        entries.append((1, f"5.{index} {screen_name}", screen_target))
        for sub_index, label in enumerate(
            ["界面设计", "字段清单", "按钮功能", "业务规则", "界面原型"], 1
        ):
            heading = f"5.{index}.{sub_index} {label}"
            sub_target = f"toc_5_{index}_{sub_index}"
            for paragraph in doc.paragraphs:
                if paragraph.text.strip() == heading:
                    _add_bookmark(paragraph, sub_target, bookmark_id)
                    bookmark_id += 1
                    break
            entries.append((2, heading, sub_target))
    entries.extend([
        (0, "6 未结与已结问题", "toc_6"),
        (1, "6.1 未结问题", "toc_6_1"),
        (1, "6.2 已结问题", "toc_6_2"),
    ])

    for position, (level, text_value, target) in enumerate(entries):
        paragraph = doc.add_paragraph()
        preface._element.addprevious(paragraph._element)
        paragraph.paragraph_format.left_indent = Pt(level * 18)
        if position == 0:
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = paragraph.add_run(text_value)
            run.bold = True
            run.font.size = Pt(16)
        elif target:
            _add_internal_link(paragraph, text_value, target)
        else:
            paragraph.add_run(text_value)


def _remove_toc(doc: Document) -> None:
    """移除模板中的目录内容控件及目录域。"""
    body = doc.element.body
    for sdt in list(body.iter(qn("w:sdt"))):
        instructions = "".join(
            element.text or "" for element in sdt.iter(qn("w:instrText"))
        )
        if "TOC" in instructions.upper():
            parent = sdt.getparent()
            if parent is not None:
                parent.remove(sdt)


def _remove_report_static_toc(doc: Document) -> None:
    """移除报表示例中写死的旧目录文本。"""
    start = next((p for p in doc.paragraphs if p.text.strip() == "目录"), None)
    end = next((p for p in doc.paragraphs if p.text.strip() == "概述"), None)
    if start is None or end is None:
        return
    element = start._element
    while element is not None and element is not end._element:
        next_element = element.getnext()
        element.getparent().remove(element)
        element = next_element


def _add_bookmark(paragraph, name: str, bookmark_id: int) -> None:
    """在章节标题处放置内部跳转锚点。"""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), name)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    paragraph._p.insert(0, start)
    paragraph._p.append(end)


def _add_internal_link(paragraph, text: str, anchor: str) -> None:
    """添加无需外部关系的 Word 文档内部超链接。"""
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("w:anchor"), anchor)
    hyperlink.set(qn("w:history"), "1")
    run = OxmlElement("w:r")
    run_properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    run_properties.append(color)
    run_properties.append(underline)
    run.append(run_properties)
    text_element = OxmlElement("w:t")
    text_element.text = text
    run.append(text_element)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _mark_toc_dirty(doc: Document) -> None:
    """仅标记 TOC 域需要重算，不触发 Word 更新文档中的所有域。"""
    for instruction in doc.element.body.iter(qn("w:instrText")):
        if "TOC" not in (instruction.text or "").upper():
            continue

        paragraph = instruction
        while paragraph is not None and paragraph.tag != qn("w:p"):
            paragraph = paragraph.getparent()
        if paragraph is None:
            continue

        begin_fields = [
            element for element in paragraph.iter(qn("w:fldChar"))
            if element.get(qn("w:fldCharType")) == "begin"
        ]
        if begin_fields:
            begin_fields[-1].set(qn("w:dirty"), "true")


def _fill_report_content(doc: Document, doc_info: Dict, overview: Dict, description: Dict,
                         tech_spec: Dict, parameters: List, reports: List, test_points: List) -> None:
    """填充报表内容到 Word 文档"""
    try:
        # 找到报表章节位置
        section_index = -1
        for i, para in enumerate(doc.paragraphs):
            if "详细描述" in para.text or "技术规格" in para.text:
                section_index = i
                break

        if section_index == -1:
            section_index = len(doc.paragraphs) - 1

        insert_pos = section_index + 1

        # 填充概述
        if overview:
            p = doc.paragraphs[insert_pos] if insert_pos < len(doc.paragraphs) else doc.add_paragraph()
            p.clear()
            run = p.add_run("2. 概述")
            run.font.size = Pt(16)
            run.font.bold = True
            insert_pos += 1

            p = doc.add_paragraph()
            p.add_run(f"目的：{overview.get('purpose', '')}")
            insert_pos += 1

            p = doc.add_paragraph()
            p.add_run(f"适用范围：{overview.get('scope', '')}")
            insert_pos += 1

        # 填充详细描述
        if description:
            p = doc.add_paragraph()
            run = p.add_run("3. 详细描述")
            run.font.size = Pt(16)
            run.font.bold = True
            insert_pos += 1

            p = doc.add_paragraph()
            p.add_run(description.get('requirements', ''))
            insert_pos += 1

            if description.get('features'):
                p = doc.add_paragraph()
                run = p.add_run("主要特性：")
                run.font.bold = True
                insert_pos += 1

                for feature in description.get('features', []):
                    p = doc.add_paragraph()
                    p.add_run(f"• {feature}")
                    insert_pos += 1

        # 填充技术规格
        if tech_spec:
            p = doc.add_paragraph()
            run = p.add_run("4. 技术规格")
            run.font.size = Pt(16)
            run.font.bold = True
            insert_pos += 1

            p = doc.add_paragraph()
            p.add_run(f"参考报表：{tech_spec.get('reference_report', '')}")
            insert_pos += 1

            p = doc.add_paragraph()
            p.add_run(f"数据范围：{tech_spec.get('data_scope', '')}")
            insert_pos += 1

            # 添加字段表格
            if tech_spec.get('fields'):
                p = doc.add_paragraph()
                run = p.add_run("字段说明：")
                run.font.bold = True
                insert_pos += 1

                table = doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'

                header_cells = table.rows[0].cells
                header_cells[0].text = '序号'
                header_cells[1].text = '字段名'
                header_cells[2].text = '类型'
                header_cells[3].text = '来源'
                header_cells[4].text = '说明'

                for field in tech_spec.get('fields', []):
                    row = table.add_row().cells
                    row[0].text = str(field.get('seq', ''))
                    row[1].text = field.get('name', '')
                    row[2].text = field.get('type', '')
                    row[3].text = field.get('source', '')
                    row[4].text = field.get('desc', '')

                insert_pos += 1

        # 填充报表参数
        if parameters:
            p = doc.add_paragraph()
            run = p.add_run("6. 报表参数")
            run.font.size = Pt(16)
            run.font.bold = True
            insert_pos += 1

            table = doc.add_table(rows=1, cols=5)
            table.style = 'Table Grid'

            header_cells = table.rows[0].cells
            header_cells[0].text = '参数名'
            header_cells[1].text = '类型'
            header_cells[2].text = '必填'
            header_cells[3].text = '默认值'
            header_cells[4].text = '说明'

            for param in parameters:
                row = table.add_row().cells
                row[0].text = param.get('name', '')
                row[1].text = param.get('type', '')
                row[2].text = param.get('required', '')
                row[3].text = param.get('default', '')
                row[4].text = param.get('desc', '')

            insert_pos += 1

        # 填充测试要点
        if test_points:
            p = doc.add_paragraph()
            run = p.add_run("8. 测试要点")
            run.font.size = Pt(16)
            run.font.bold = True
            insert_pos += 1

            for i, point in enumerate(test_points, 1):
                p = doc.add_paragraph()
                p.add_run(f"{i}. {point}")
                insert_pos += 1

        logger.info(f"成功填充报表内容，{len(parameters)} 个参数，{len(test_points)} 个测试要点")

    except Exception as e:
        logger.error(f"填充报表内容失败: {e}", exc_info=True)


# ==================== 历史文档管理 ====================


@apex_ai_bp.route("/exports", methods=["GET"])
def list_exports():
    """列出已生成的历史文档"""
    try:
        items = []
        for file_path in sorted(EXPORTS_DIR.glob("*.docx"), key=lambda p: p.stat().st_mtime, reverse=True):
            stat = file_path.stat()
            # 从文件名提取标题（去掉时间戳）
            filename = file_path.name
            title = filename.rsplit('_', 2)[0] if '_' in filename else filename.replace('.docx', '')

            # 检查是否有对应的 JSON 预览文件
            json_filename = filename.replace('.docx', '.json')
            json_path = EXPORTS_DIR / json_filename
            has_preview = json_path.exists()
            template_key = "form"
            if has_preview:
                try:
                    with open(json_path, "r", encoding="utf-8") as f:
                        preview_data = json.load(f)
                    template_key = preview_data.get("template_key", "form")
                except (OSError, json.JSONDecodeError) as exc:
                    logger.warning(f"读取历史文档类型失败: {json_path.name}, {exc}")

            template_label = "报表类" if template_key == "report" else "表单类"

            items.append({
                "title": title,
                "filename": filename,
                "download_url": f"/apex_ai/exports/{filename}",
                "generated_at": datetime.fromtimestamp(stat.st_mtime).strftime('%Y-%m-%d %H:%M:%S'),
                "size": stat.st_size,
                "template_key": template_key,
                "template_label": template_label,
                "has_preview": has_preview
            })

        return jsonify({"success": True, "items": items})

    except Exception as e:
        logger.error(f"列出历史文档失败: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


@apex_ai_bp.route("/exports/<filename>", methods=["GET"])
def download_export(filename: str):
    """下载已生成的文档"""
    try:
        file_path = EXPORTS_DIR / filename
        if not file_path.exists():
            return jsonify({"success": False, "error": "文件不存在"}), 404

        return send_file(
            str(file_path),
            as_attachment=True,
            download_name=filename,
            mimetype="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

    except Exception as e:
        logger.error(f"下载文档失败: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500


@apex_ai_bp.route("/preview-history/<filename>", methods=["GET"])
def preview_history_document(filename: str):
    """预览历史文档（从保存的 JSON 数据生成 HTML）"""
    try:
        # 查找对应的 JSON 文件
        json_filename = filename.replace('.docx', '.json')
        json_path = EXPORTS_DIR / json_filename

        if not json_path.exists():
            return jsonify({"success": False, "error": "预览数据不存在"}), 404

        # 读取 JSON 数据
        with open(json_path, 'r', encoding='utf-8') as f:
            preview_data = json.load(f)

        template_key = preview_data.get("template_key", "form")
        data = preview_data.get("data", {})

        # 生成 HTML 预览
        html = _generate_report_preview(data) if template_key == "report" else _generate_form_preview(data)

        return jsonify({"success": True, "html": html, "data": data, "template_key": template_key})

    except Exception as e:
        logger.error(f"预览历史文档失败: {e}", exc_info=True)
        return jsonify({"success": False, "error": str(e)}), 500
    pass
